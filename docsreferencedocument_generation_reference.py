import os
import shutil
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime

try:
    from database import (
        get_connection,
        meetings_add,
        meetings_get_all,
        meetings_update,
        meetings_delete,
        get_all_employees,
        calendar_add_task,
        invited_get_all
    )
except ImportError as e:
    messagebox.showerror("Ошибка", f"Не удалось импортировать функции базы данных:\n{e}")
    raise

def apply_modern_style(window):
    style = ttk.Style()

    try:
        style.theme_use("clam")
    except tk.TclError:
        pass

    style.configure(
        "Modern.TButton",
        font=("Segoe UI", 12),
        padding=(10, 6),
        background="#1f3b63",
        foreground="white"
    )

    style.map(
        "Modern.TButton",
        background=[("active", "#444444")],
        foreground=[("active", "white")]
    )

    style.configure(
        "Gray.TButton",
        font=("Segoe UI", 12, "bold"),
        padding=(10, 6),
        background="#777777",
        foreground="white"
    )

    style.map(
        "Gray.TButton",
        background=[("active", "#5f5f5f")],
        foreground=[("active", "white")]
    )

    style.configure(
        "Sidebar.TButton",
        font=("Segoe UI", 12),
        padding=(10, 8),
        background="#1f3b63",
        foreground="white",
        relief="flat"
    )

    style.map(
        "Sidebar.TButton",
        background=[("active", "#505050")]
    )

    style.configure("Modern.TEntry", fieldbackground="white")

    style.configure(
        "Selected.TButton",
        font=("Segoe UI", 12, "bold"),
        padding=(10, 6),
        background="#b7d9f3",
        foreground="#0f2f4a"
    )

    style.map(
        "Selected.TButton",
        background=[("active", "#9fc9e8")],
        foreground=[("active", "#0f2f4a")]
    )


class MultiLineEntry(tk.Text):
    """
    Небольшая обертка над tk.Text, чтобы поле "По вопросу"
    могло быть многострочным, но старый код мог работать с ним
    как с обычным Entry: get(), delete(0, "end"), insert(0, text).
    """
    def __init__(self, master=None, **kwargs):
        kwargs.setdefault("height", 2)
        kwargs.setdefault("wrap", "word")
        kwargs.setdefault("relief", "solid")
        kwargs.setdefault("bd", 1)
        super().__init__(master, **kwargs)

    def get(self, *args):
        if not args:
            return super().get("1.0", "end").strip()
        return super().get(*args)

    def delete(self, *args):
        if args and args[0] == 0:
            return super().delete("1.0", "end")
        return super().delete(*args)

    def insert(self, *args):
        if args and args[0] == 0:
            return super().insert("1.0", args[1] if len(args) > 1 else "")
        return super().insert(*args)


class SelectParticipantsWindow(tk.Toplevel):
    def __init__(self, parent, current_list_str, target="local"):
        super().__init__(parent)
        apply_modern_style(self)
        self.configure(bg="#f4f4f4")
        self.title("Выберите участников")
        self.geometry("1000x650")
        self.transient(parent)
        self.grab_set()
        self.target = target

        current_names = [name.strip() for name in current_list_str.split('\n') if name.strip()]

        try:
            if target == "invited":
                all_people = self._load_invited_people_records()

                if not all_people:
                    all_people = invited_get_all()
            else:
                all_people = self._load_local_people_records()

                if not all_people:
                    all_people = get_all_employees()
        except Exception as e:
            messagebox.showerror("Ошибка БД", f"Не удалось загрузить список:\n{e}")
            self.destroy()
            return

        if not all_people:
            messagebox.showinfo("Инфо", "Список пуст.")
            self.destroy()
            return

        title_text = "Отметьте приглашенных:" if target == "invited" else "Отметьте участников совещания:"
        ttk.Label(self, text=title_text, font=("Segoe UI", 10, "bold")).pack(pady=10)

        scroll_frame = tk.Frame(self)
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.canvas = tk.Canvas(scroll_frame, borderwidth=0, background="#ffffff")
        self.list_frame = tk.Frame(self.canvas, background="#ffffff")

        vsb = ttk.Scrollbar(scroll_frame, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=vsb.set)

        vsb.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        self.canvas_window = self.canvas.create_window(
            (4, 4),
            window=self.list_frame,
            anchor="nw",
            tags="list_frame"
        )

        def _on_mousewheel(event):
            # Windows / macOS
            if event.delta:
                self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            # Linux
            elif event.num == 4:
                self.canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                self.canvas.yview_scroll(1, "units")

        def _bind_mousewheel(event=None):
            self.canvas.bind_all("<MouseWheel>", _on_mousewheel)
            self.canvas.bind_all("<Button-4>", _on_mousewheel)
            self.canvas.bind_all("<Button-5>", _on_mousewheel)

        def _unbind_mousewheel(event=None):
            self.canvas.unbind_all("<MouseWheel>")
            self.canvas.unbind_all("<Button-4>")
            self.canvas.unbind_all("<Button-5>")

        self.canvas.bind("<Enter>", _bind_mousewheel)
        self.canvas.bind("<Leave>", _unbind_mousewheel)
        self.list_frame.bind("<Enter>", _bind_mousewheel)
        self.list_frame.bind("<Leave>", _unbind_mousewheel)


        def _resize_list_frame(event):
            self.canvas.itemconfig(self.canvas_window, width=event.width)

        self.canvas.bind("<Configure>", _resize_list_frame)

        self.check_vars = {}
        self.expanded_sections = {}
        self.triangle_labels = {}
        self.section_frames = {}

        if target == "invited":
            self._create_invited_list(all_people, current_names)
        else:
            self._create_local_list(all_people, current_names)

        self._bind_mousewheel_to_children(self.list_frame)
        self.list_frame.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

        btn_frame = tk.Frame(self, bg="#f0f0f0")
        btn_frame.pack(fill="x", padx=10, pady=10)

        ttk.Button(btn_frame, text="✅ Добавить выбранных", command=self.on_confirm).pack(side="right", padx=5)
        ttk.Button(btn_frame, text="❌ Отмена", command=self.destroy).pack(side="right", padx=5)

    def _person_name(self, person):
        return (
            person.get("full_name")
            or person.get("fio")
            or person.get("name")
            or ""
        ).strip()

    def _person_position(self, person):
        return (
            person.get("position")
            or person.get("dolzhnost")
            or person.get("job_title")
            or ""
        ).strip()

    def _person_leadership(self, person):
        return (
            person.get("leadership")
            or person.get("rukovodstvo")
            or person.get("supervisor")
            or person.get("boss")
            or ""
        ).strip()

    def _normalize_text(self, text):
        text = str(text or "").strip().lower().replace("ё", "е")
        return " ".join(text.split())

    def _load_invited_people_records(self):
        """
        Загружает приглашенных из новой таблицы meeting_participants,
        которая заполняется в настройках: Участники совещаний -> Список приглашенных.
        Если таблица пустая или старой структуры, вызывающий код использует старый invited_get_all().
        """
        try:
            with get_connection() as conn:
                cursor = conn.cursor()

                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS meeting_participants (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        category TEXT NOT NULL,
                        full_name TEXT NOT NULL,
                        position TEXT DEFAULT '',
                        leadership TEXT DEFAULT '',
                        sort_order INTEGER DEFAULT 999
                    )
                    """
                )

                cursor.execute(
                    """
                    SELECT id, full_name, position, leadership, sort_order
                    FROM meeting_participants
                    WHERE category = 'invited_ip'
                    ORDER BY sort_order, full_name
                    """
                )

                rows = cursor.fetchall()
                result = []

                for row in rows:
                    result.append({
                        "id": row["id"] if hasattr(row, "keys") else row[0],
                        "full_name": row["full_name"] if hasattr(row, "keys") else row[1],
                        "position": row["position"] if hasattr(row, "keys") else row[2],
                        "leadership": row["leadership"] if hasattr(row, "keys") else row[3],
                        "sort_order": row["sort_order"] if hasattr(row, "keys") else row[4],
                    })

                return result

        except Exception:
            return []


    def _load_local_people_records(self):
        """
        Загружает участников из новой таблицы meeting_participants,
        которая заполняется в настройках: Участники совещаний -> Список органов МСУ.

        Поэтому все, что добавлено в настройках в "Список органов МСУ",
        автоматически появляется при нажатии кнопки "Органы МСУ" в окне совещаний.
        """
        try:
            with get_connection() as conn:
                cursor = conn.cursor()

                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS meeting_participants (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        category TEXT NOT NULL,
                        full_name TEXT NOT NULL,
                        position TEXT DEFAULT '',
                        leadership TEXT DEFAULT '',
                        sort_order INTEGER DEFAULT 999
                    )
                    """
                )

                cursor.execute(
                    """
                    SELECT id, full_name, position, leadership, sort_order
                    FROM meeting_participants
                    WHERE category = 'msu_ip'
                    ORDER BY sort_order, full_name
                    """
                )

                rows = cursor.fetchall()
                result = []

                for row in rows:
                    result.append({
                        "id": row["id"] if hasattr(row, "keys") else row[0],
                        "full_name": row["full_name"] if hasattr(row, "keys") else row[1],
                        "position": row["position"] if hasattr(row, "keys") else row[2],
                        "leadership": row["leadership"] if hasattr(row, "keys") else row[3],
                        "sort_order": row["sort_order"] if hasattr(row, "keys") else row[4],
                        "source_table": "meeting_participants",
                    })

                return result

        except Exception:
            return []

    def _bind_mousewheel_to_children(self, widget):
        for child in widget.winfo_children():
            child.bind("<MouseWheel>", lambda e:self.canvas.yview_scroll(int(-1 * (e.delta / 120)), "units"))
            child.bind("<Button-4>", lambda e:self.canvas.yview_scroll(-1, "units"))
            child.bind("<Button-5>", lambda e:self.canvas.yview_scroll(1, "units"))
            self._bind_mousewheel_to_children(child)


    def _create_invited_list(self, invited_people, current_names):
        """
        Отображает приглашенных с такой же визуальной подчиненностью,
        как список "Органы МСУ": руководитель сверху, подчиненные внутри
        раскрывающегося блока. Связь берется из поля leadership, которое
        заполняется в настройках: Участники совещаний -> Список приглашенных.
        """
        invited_people = list(invited_people or [])

        if not invited_people:
            return

        def sort_key(person):
            return (
                self._normalize_text(self._person_position(person)),
                self._normalize_text(self._person_name(person))
            )

        children_by_leadership = {}
        roots = []

        for person in invited_people:
            name = self._person_name(person)

            if not name:
                continue

            leadership = self._normalize_text(self._person_leadership(person))

            if leadership:
                children_by_leadership.setdefault(leadership, []).append(person)
            else:
                roots.append(person)

        for key in list(children_by_leadership.keys()):
            children_by_leadership[key] = sorted(children_by_leadership[key], key=sort_key)

        roots = sorted(roots, key=sort_key)
        used_ids = set()

        def person_key(person):
            return str(person.get("id") or self._person_name(person))

        def create_invited_section(person, level=0):
            name = self._person_name(person)
            position = self._person_position(person)

            if not name:
                return

            used_ids.add(person_key(person))

            section_frame = tk.Frame(self.list_frame, bg="#ffffff")
            section_frame.pack(fill="x", pady=2, padx=5 + level * 18)

            header_frame = tk.Frame(section_frame, bg="#1976d2")
            header_frame.pack(fill="x")

            child_key = self._normalize_text(position)
            children = children_by_leadership.get(child_key, [])

            triangle_lbl = tk.Label(
                header_frame,
                text="▼" if children else "",
                bg="#1976d2",
                fg="white",
                font=("Segoe UI", 10, "bold"),
                width=3,
                cursor="hand2" if children else "arrow"
            )
            triangle_lbl.pack(side="left", padx=5)

            label_text = name
            if position:
                label_text += f" ({position})"

            lbl = tk.Label(
                header_frame,
                text=label_text,
                font=("Segoe UI", 10, "bold"),
                bg="#1976d2",
                fg="white",
                anchor="w",
                padx=10,
                pady=5
            )
            lbl.pack(side="left", fill="x", expand=True)

            var = tk.BooleanVar(value=(name in current_names))
            self.check_vars[name] = var

            tk.Checkbutton(
                header_frame,
                variable=var,
                bg="#1976d2",
                activebackground="#1976d2",
                selectcolor="#1976d2"
            ).pack(side="right", padx=10)

            sub_frame = tk.Frame(section_frame, bg="#f5f5f5")

            if children:
                sub_frame.pack(fill="x", padx=5, pady=2)

                self.section_frames[name] = sub_frame
                self.expanded_sections[name] = True
                self.triangle_labels[name] = triangle_lbl

                def toggle(event=None, section_name=name):
                    self._toggle_section(section_name)

                triangle_lbl.bind("<Button-1>", toggle)
                lbl.bind("<Button-1>", toggle)

                for child in children:
                    if person_key(child) in used_ids:
                        continue

                    child_name = self._person_name(child)
                    child_position = self._person_position(child)

                    row = tk.Frame(sub_frame, bg="#f5f5f5")
                    row.pack(fill="x", padx=10, pady=3)

                    text = f"● {child_name}"
                    if child_position:
                        text += f" — {child_position}"

                    child_var = tk.BooleanVar(value=(child_name in current_names))
                    self.check_vars[child_name] = child_var
                    used_ids.add(person_key(child))

                    tk.Checkbutton(
                        row,
                        text=text,
                        variable=child_var,
                        font=("Segoe UI", 10),
                        bg="#f5f5f5",
                        anchor="w",
                        justify="left"
                    ).pack(fill="x", padx=5, pady=2)

        for root in roots:
            create_invited_section(root)

        remaining = [
            person for person in invited_people
            if person_key(person) not in used_ids and self._person_name(person)
        ]

        for person in sorted(remaining, key=sort_key):
            create_invited_section(person, level=1)

    def _create_local_list(self, all_employees, current_names):
        """
        Отображает "Органы МСУ" из таблицы meeting_participants/category='msu_ip',
        то есть из настроек "Участники совещаний -> Список органов МСУ".

        Если по какой-то причине пришел старый список get_all_employees()
        с полями level/supervisor_id, оставлен старый способ отображения.
        """
        all_employees = list(all_employees or [])

        if not all_employees:
            return

        # Новая структура из настроек.
        if any("leadership" in emp or emp.get("source_table") == "meeting_participants" for emp in all_employees):
            self._create_hierarchy_list_from_positions(all_employees, current_names)
            return

        # Старый формат из get_all_employees().
        deputies = {}

        for emp in all_employees:
            if emp.get('level') == 'deputy':
                deputies[emp['id']] = {
                    'info': emp,
                    'subordinates': []
                }

        for emp in all_employees:
            if emp.get('level') == 'head':
                sup_id = emp.get('supervisor_id')
                if sup_id and sup_id in deputies:
                    deputies[sup_id]['subordinates'].append(emp)

        sorted_deputies = sorted(
            deputies.values(),
            key=lambda x: x["info"].get("full_name", "").lower()
        )

        for dep_data in sorted_deputies:
            subordinates = sorted(
                dep_data['subordinates'],
                key=lambda x: x.get("full_name", "").lower()
            )
            self._create_deputy_section(dep_data['info'], subordinates, current_names)

    def _create_hierarchy_list_from_positions(self, people, current_names):
        """
        Иерархия для "Органы МСУ" по полю leadership из настроек.

        Важно: связь строится рекурсивно, а не только в один уровень:
        Глава города -> заместители / главы районов -> их подчиненные -> следующие уровни.
        Поэтому у каждого заместителя снова появляется свой раскрывающийся блок
        с подчиненными, как было в старом списке "Органы МСУ".
        """
        people = list(people or [])

        def sort_key(person):
            return (
                self._normalize_text(self._person_position(person)),
                self._normalize_text(self._person_name(person))
            )

        def person_key(person):
            return str(person.get("id") or self._person_name(person))

        children_by_leadership = {}
        roots = []

        for person in people:
            name = self._person_name(person)

            if not name:
                continue

            leadership = self._normalize_text(self._person_leadership(person))

            if leadership:
                children_by_leadership.setdefault(leadership, []).append(person)
            else:
                roots.append(person)

        for key in list(children_by_leadership.keys()):
            children_by_leadership[key] = sorted(children_by_leadership[key], key=sort_key)

        roots = sorted(roots, key=sort_key)
        used_ids = set()

        def create_plain_row(person, parent_frame, level=0):
            child_name = self._person_name(person)
            child_position = self._person_position(person)

            if not child_name:
                return

            row = tk.Frame(parent_frame, bg="#f5f5f5")
            row.pack(fill="x", padx=10 + level * 14, pady=3)

            text = f"● {child_name}"
            if child_position:
                text += f" — {child_position}"

            child_var = tk.BooleanVar(value=(child_name in current_names))
            self.check_vars[child_name] = child_var

            tk.Checkbutton(
                row,
                text=text,
                variable=child_var,
                font=("Segoe UI", 10),
                bg="#f5f5f5",
                anchor="w",
                justify="left"
            ).pack(fill="x", padx=5, pady=2)

        def create_section(person, parent_frame=None, level=0):
            name = self._person_name(person)
            position = self._person_position(person)

            if not name:
                return

            current_key = person_key(person)
            if current_key in used_ids:
                return

            used_ids.add(current_key)

            child_key = self._normalize_text(position)
            children = [
                child for child in children_by_leadership.get(child_key, [])
                if person_key(child) not in used_ids
            ]

            # Для отдельных руководителей оставляем такое же оформление,
            # как у записей с подчиненными, даже если в текущей базе их подчиненные
            # временно не подтянулись.
            force_header_names = (
                "данькин",
                "асеев",
                "андреева",
            )
            force_header = any(mark in self._normalize_text(name) for mark in force_header_names)

            # Если это подчиненный без своих подчиненных, показываем его обычной строкой
            # внутри блока руководителя. Если это корневая запись или принудительно
            # оформляемый руководитель, оставляем синий заголовок с галочкой.
            if parent_frame is not None and not children and not force_header:
                create_plain_row(person, parent_frame, level=0)
                return

            container = parent_frame if parent_frame is not None else self.list_frame

            section_frame = tk.Frame(container, bg="#ffffff" if parent_frame is None else "#f5f5f5")
            section_frame.pack(fill="x", pady=2, padx=5 + (level * 18 if parent_frame is None else 10))

            header_frame = tk.Frame(section_frame, bg="#1976d2")
            header_frame.pack(fill="x")

            triangle_lbl = tk.Label(
                header_frame,
                text="▼" if children else "",
                bg="#1976d2",
                fg="white",
                font=("Segoe UI", 10, "bold"),
                width=3,
                cursor="hand2" if children else "arrow"
            )
            triangle_lbl.pack(side="left", padx=5)

            label_text = name
            if position:
                label_text += f" ({position})"

            lbl = tk.Label(
                header_frame,
                text=label_text,
                font=("Segoe UI", 10, "bold"),
                bg="#1976d2",
                fg="white",
                anchor="w",
                padx=10,
                pady=5
            )
            lbl.pack(side="left", fill="x", expand=True)

            var = tk.BooleanVar(value=(name in current_names))
            self.check_vars[name] = var

            tk.Checkbutton(
                header_frame,
                variable=var,
                bg="#1976d2",
                activebackground="#1976d2",
                selectcolor="#1976d2"
            ).pack(side="right", padx=10)

            sub_frame = tk.Frame(section_frame, bg="#f5f5f5")

            if children:
                sub_frame.pack(fill="x", padx=5, pady=2)

                self.section_frames[name] = sub_frame
                self.expanded_sections[name] = True
                self.triangle_labels[name] = triangle_lbl

                def toggle(event=None, section_name=name):
                    self._toggle_section(section_name)

                triangle_lbl.bind("<Button-1>", toggle)
                lbl.bind("<Button-1>", toggle)

                for child in children:
                    create_section(child, sub_frame, level + 1)

        for root in roots:
            create_section(root, None, 0)

        # Если у записи указано руководство, но соответствующий руководитель
        # не найден в списке, показываем ее отдельным блоком внизу, чтобы она
        # не пропала из окна выбора.
        remaining = [
            person for person in people
            if person_key(person) not in used_ids and self._person_name(person)
        ]

        for person in sorted(remaining, key=sort_key):
            create_section(person, None, 1)

    def _create_deputy_section(self, deputy, subordinates, current_names):
        dep_name = deputy['full_name']
        dep_pos = deputy.get('position', '')

        deputy_frame = tk.Frame(self.list_frame, bg="#ffffff")
        deputy_frame.pack(fill="x", pady=2, padx=5)

        header_frame = tk.Frame(deputy_frame, bg="#1976d2")
        header_frame.pack(fill="x")

        triangle_lbl = tk.Label(
            header_frame,
            text="▼",
            bg="#1976d2",
            fg="white",
            font=("Segoe UI", 10, "bold"),
            width=3,
            cursor="hand2"
        )
        triangle_lbl.pack(side="left", padx=5)
        self.triangle_labels[dep_name] = triangle_lbl

        lbl = tk.Label(
            header_frame,
            text=f"{dep_name} ({dep_pos})",
            font=("Segoe UI", 10, "bold"),
            bg="#1976d2",
            fg="white",
            anchor="w",
            padx=10,
            pady=5
        )
        lbl.pack(side="left", fill="x", expand=True)

        var = tk.BooleanVar(value=(dep_name in current_names))
        self.check_vars[dep_name] = var

        tk.Checkbutton(
            header_frame,
            variable=var,
            bg="#1976d2",
            activebackground="#1976d2",
            selectcolor="#1976d2"
        ).pack(side="right", padx=10)

        def toggle(event, name=dep_name):
            self._toggle_section(name)

        triangle_lbl.bind("<Button-1>", toggle)
        lbl.bind("<Button-1>", toggle)

        sub_frame = tk.Frame(deputy_frame, bg="#f5f5f5")
        sub_frame.pack(fill="x", padx=5, pady=2)

        self.section_frames[dep_name] = sub_frame
        self.expanded_sections[dep_name] = True

        if subordinates:
            for sub in subordinates:
                self._create_subordinate_section(sub, sub_frame, current_names)
        else:
            tk.Label(
                sub_frame,
                text=" Нет подчиненных",
                font=("Segoe UI", 10),
                bg="#f5f5f5",
                fg="#999"
            ).pack(fill="x", padx=10, pady=5)

    def _create_subordinate_section(self, subordinate, parent_frame, current_names):
        sub_name = subordinate['full_name']
        dept = subordinate.get('department', '')

        sub_frame = tk.Frame(parent_frame, bg="#f5f5f5")
        sub_frame.pack(fill="x", padx=10, pady=3)

        tk.Label(
            sub_frame,
            text=f"● {sub_name}",
            font=("Segoe UI", 10, "bold"),
            bg="#f5f5f5",
            anchor="w",
            padx=5
        ).pack(fill="x")

        if dept:
            tk.Label(
                sub_frame,
                text=f" {dept}",
                font=("Segoe UI", 10),
                bg="#f5f5f5",
                fg="#555",
                anchor="w",
                padx=5
            ).pack(fill="x")

        var = tk.BooleanVar(value=(sub_name in current_names))
        self.check_vars[sub_name] = var

        tk.Checkbutton(
            sub_frame,
            text=" Выбрать",
            variable=var,
            font=("Segoe UI", 10),
            bg="#f5f5f5",
            anchor="w"
        ).pack(fill="x", pady=2, padx=10)

    def _toggle_section(self, deputy_name):
        sf = self.section_frames.get(deputy_name)
        tl = self.triangle_labels.get(deputy_name)

        if sf and tl:
            is_vis = self.expanded_sections.get(deputy_name, False)

            if is_vis:
                sf.pack_forget()
                tl.config(text="▶")
                self.expanded_sections[deputy_name] = False
            else:
                sf.pack(fill="x", padx=5, pady=2)
                tl.config(text="▼")
                self.expanded_sections[deputy_name] = True

            self.list_frame.update_idletasks()
            self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def on_confirm(self):
        result = [name for name, var in self.check_vars.items() if var.get()]

        if self.target == "invited":
            if hasattr(self.master, "on_invited_selected"):
                self.master.on_invited_selected(result)
        else:
            if hasattr(self.master, "on_participants_selected"):
                self.master.on_participants_selected(result)

        self.destroy()


class MeetingsWindow(tk.Toplevel):
    def __init__(self, parent, current_user=None, initial_doc_type=None):
        super().__init__(parent)
        # Тип документа выбирается ДО создания основного окна в show_meetings().
        # Поэтому окно больше не скрывается через withdraw(): из-за скрытого родителя
        # модальное окно выбора могло попасть за другие окна и программа выглядела зависшей.
        self._initial_doc_type_chosen = bool(initial_doc_type)
        self.current_user = current_user or "Администратор"
        apply_modern_style(self)
        self.configure(bg="#f4f6f9")
        self.parent_main = parent
        self.title("Управление совещаниями")
        self.geometry("1200x850")
        self.minsize(1180, 820)
        self.resizable(True, True)
        # Окно совещаний теперь открывается чуть меньше, без режима "на весь экран".

        self.current_id = None
        self.attachment_var = tk.StringVar(value="Файл не прикреплен")
        self.current_file_path = ""
        self.attached_files = []
        self.temp_selected_participants = []
        self.temp_invited_participants = []
        self.agenda_question_rows = []
        self.agenda_speaker_rows = []
        self.tasks_mode = "tasks"
        self.attachment_type = tk.StringVar(value="none")
        self.telegram_number = tk.StringVar(value="№ 200/05/ИТФ___")
        self.doc_type_buttons = {}
        self.transfer_fio = tk.StringVar(value="Иванова Елена Николаевна")
        self.transfer_phone = tk.StringVar(value="")
        self.transfer_email = tk.StringVar(value="fedorova-en@barnaul-adm.ru")
        self.telegram_sign_fio = tk.StringVar(value="О.А. Финк")

        # Поля для документа "Протокол"
        self.protocol_number = tk.StringVar(value="200/05/ПРОТ-___")
        self.protocol_chair_fio = tk.StringVar(value="О.А. Финк")
        self.protocol_chair_position = tk.StringVar(value="заместитель главы администрации города, руководитель аппарата")

        self.protocol_keeper = tk.StringVar(value="Иванова Елена Николаевна")

        # Кабинет для всех документов
        self.cabinet_number = tk.StringVar(value="213")


        # Поля подписи для повестки
        self.agenda_sign_position = tk.StringVar(value="Председатель правового комитета")
        self.agenda_sign_fio = tk.StringVar(value="О.И. Насыров")

        self._ensure_meetings_extra_columns()
        self.agenda_placeholder_text = "О ходе исполнения..."
        self.agenda_placeholder_active = False

        style = ttk.Style()
        style.theme_use("clam")
        style.configure("Treeview", rowheight=30, font=("Segoe UI", 10))
        style.configure("Treeview.Heading", font=("Segoe UI", 10, "bold"))
        style.configure("Header.TLabel", font=("Segoe UI", 10, "bold"), background="#f5f5f5")

        form_container = tk.Frame(self, bg="#f5f5f5")
        form_container.pack(fill="x", padx=10, pady=5)

        self.header_frame = tk.Frame(form_container, bg="#f5f5f5")
        self.header_frame.pack(fill="x", pady=(0, 2))

        ttk.Label(
            self.header_frame,
            text="Выберите тип создаваемого документа:",
            style="Header.TLabel"
        ).pack(side="left", padx=(0, 8))

        # Кнопки выбора типа документа теперь находятся в самой верхней строке окна.
        self.doc_buttons_frame = tk.Frame(self.header_frame, bg="#f5f5f5")
        self.doc_buttons_frame.pack(side="left", fill="x", expand=True)

        self.details_title_frame = tk.Frame(form_container, bg="#f5f5f5")
        self.details_title_frame.pack(fill="x", pady=(2, 5))

        ttk.Label(
            self.details_title_frame,
            text="Детали совещания",
            style="Header.TLabel"
        ).pack(anchor="w")

        # Верхний блок с полями. Номер показывается только для
        # "Телефонограммы" и "Протокола".
        self.top_telegram_number_block = tk.Frame(form_container, bg="#f5f5f5")
        ttk.Label(
            self.top_telegram_number_block,
            text="№:",
            width=12,
            anchor="e"
        ).pack(side="left", padx=5)
        self.ent_telegram_number = tk.Entry(
            self.top_telegram_number_block,
            textvariable=self.telegram_number,
            font=("Segoe UI", 10),
            width=24
        )
        self.ent_telegram_number.pack(side="left", padx=5)

        self.top_protocol_number_block = tk.Frame(form_container, bg="#f5f5f5")
        ttk.Label(
            self.top_protocol_number_block,
            text="№:",
            width=12,
            anchor="e"
        ).pack(side="left", padx=5)
        tk.Entry(
            self.top_protocol_number_block,
            textvariable=self.protocol_number,
            font=("Segoe UI", 10),
            width=24
        ).pack(side="left", padx=5)

        self.row_title = tk.Frame(form_container, bg="#f5f5f5")
        self.row_title.pack(fill="x", pady=2)

        ttk.Label(self.row_title, text="По вопросу:", width=12, anchor="e").pack(side="left", padx=5)
        self.ent_title = MultiLineEntry(self.row_title, font=("Segoe UI", 10), width=70, height=2)
        self.ent_title.pack(side="left", padx=5, fill="x", expand=True)

        self.row_date = tk.Frame(form_container, bg="#f5f5f5")
        self.row_date.pack(fill="x", pady=2)

        ttk.Label(self.row_date, text="Дата:", width=12, anchor="e").pack(side="left", padx=5)
        self.ent_date = tk.Entry(self.row_date, font=("Segoe UI", 10), width=14)
        self.ent_date.insert(0, datetime.now().strftime("%d.%m.%Y"))
        self.ent_date.pack(side="left", padx=5)

        ttk.Label(self.row_date, text="Время:", width=8, anchor="e").pack(side="left", padx=(12, 0))
        self.ent_time = tk.Entry(self.row_date, font=("Segoe UI", 10), width=12)
        self.ent_time.insert(0, datetime.now().strftime("%H:%M"))
        self.ent_time.pack(side="left", padx=5)

        ttk.Label(self.row_date, text="Кабинет:", width=9, anchor="e").pack(side="left", padx=(12, 0))
        self.ent_cabinet = tk.Entry(
            self.row_date,
            textvariable=self.cabinet_number,
            font=("Segoe UI", 10),
            width=10
        )
        self.ent_cabinet.pack(side="left", padx=5)

        # Поле "Протокол вела" для раздела "Протокол".
        # Показывается на этой же строке справа от поля "Кабинет".
        self.row_date_protocol_keeper_label = ttk.Label(
            self.row_date,
            text="Протокол вела:",
            width=15,
            anchor="e"
        )
        self.row_date_protocol_keeper_entry = tk.Entry(
            self.row_date,
            textvariable=self.protocol_keeper,
            font=("Segoe UI", 10),
            width=26
        )

        self.row2 = tk.Frame(form_container, bg="#f5f5f5")
        self.row2.pack(fill="both", expand=True, pady=5)

        self.left_col = tk.Frame(self.row2, bg="#f5f5f5")
        self.left_col.pack(side="left", fill="both", expand=True, padx=(0, 5))

        left_col = self.left_col

        # ===== ВОПРОС В ПОВЕСТКЕ =====
        self.agenda_label = ttk.Label(left_col, text="", style="Header.TLabel")
        self.agenda_label.pack(anchor="w")

        self.agenda_block = tk.Frame(left_col, bg="#f5f5f5", height=90)
        self.agenda_block.pack(fill="x", expand=False, pady=(2, 0))
        self.agenda_block.pack_propagate(False)

        agenda_outer = tk.Frame(self.agenda_block, bg="#ffffff", relief="solid", bd=1)
        agenda_outer.pack(fill="both", expand=True)

        self.agenda_canvas = tk.Canvas(
            agenda_outer,
            bg="white",
            highlightthickness=0,
            height=80
        )

        self.agenda_scroll = ttk.Scrollbar(
            agenda_outer,
            orient="vertical",
            command=self.agenda_canvas.yview
        )

        self.agenda_container = tk.Frame(self.agenda_canvas, bg="white")

        self.agenda_container.bind(
            "<Configure>",
            lambda e:self.agenda_canvas.configure(scrollregion=self.agenda_canvas.bbox("all"))
        )

        self.agenda_canvas_window = self.agenda_canvas.create_window(
            (0, 0),
            window=self.agenda_container,
            anchor="nw"
        )

        def resize_agenda_container(event):
            self.agenda_canvas.itemconfig(self.agenda_canvas_window, width=event.width)

        self.agenda_canvas.bind("<Configure>", resize_agenda_container)
        self.agenda_canvas.configure(yscrollcommand=self.agenda_scroll.set)

        self.agenda_canvas.pack(side="left", fill="both", expand=True)
        self.agenda_scroll.pack(side="right", fill="y")

        self.agenda_question_rows = []

        def render_agenda_questions():
            for widget in self.agenda_container.winfo_children():
                widget.destroy()

            for i, row_data in enumerate(self.agenda_question_rows, start=1):
                row = tk.Frame(self.agenda_container, bg="white")
                row.pack(fill="x", pady=2, padx=5)

                row.grid_columnconfigure(1, weight=1)

                tk.Label(
                    row,
                    text=f"{i}.",
                    width=3,
                    bg="white"
                ).grid(row=0, column=0, padx=(0, 5), sticky="w")

                entry = tk.Entry(
                    row,
                    textvariable=row_data["text"],
                    font=("Segoe UI", 10)
                )
                entry.grid(row=0, column=1, padx=5, sticky="ew")

            self.agenda_container.update_idletasks()
            self.agenda_canvas.configure(scrollregion=self.agenda_canvas.bbox("all"))

        def add_agenda_question_row(text=""):
            var = tk.StringVar(value=text)
            self.agenda_question_rows.append({"text":var})
            render_agenda_questions()
            self.agenda_canvas.yview_moveto(1.0)

        def remove_agenda_question_row():
            if self.agenda_question_rows:
                self.agenda_question_rows.pop()
                render_agenda_questions()

        def clear_agenda_question_rows():
            self.agenda_question_rows.clear()
            add_agenda_question_row("")

        self.render_agenda_questions = render_agenda_questions
        self.add_agenda_question_row = add_agenda_question_row
        self.remove_agenda_question_row = remove_agenda_question_row
        self.clear_agenda_question_rows = clear_agenda_question_rows

        agenda_btns = tk.Frame(left_col, bg="#f5f5f5")
        agenda_btns.pack(fill="x")

        self.agenda_add_btn = tk.Button(
            agenda_btns,
            text="+",
            command=self.add_agenda_question_row
        )
        self.agenda_add_btn.pack(side="left", padx=5, pady=3)

        self.agenda_remove_btn = tk.Button(
            agenda_btns,
            text="-",
            command=self.remove_agenda_question_row
        )
        self.agenda_remove_btn.pack(side="left", padx=5, pady=3)

        # Оставляем скрытое поле для совместимости со старым кодом,
        # чтобы не падали старые обращения к self.txt_agenda.
        self.txt_agenda = tk.Text(self, height=1)
        self.txt_agenda.pack_forget()

        add_agenda_question_row("")

        self.protocol_label = ttk.Label(left_col, text="Протокол/Задачи:", style="Header.TLabel")
        self.protocol_label.pack(anchor="w", pady=(10, 2))

        self.protocol_block = tk.Frame(left_col, bg="#f5f5f5")
        self.protocol_block.pack(fill="both", expand=True, pady=(2, 0))

        self.txt_protocol = tk.Text(
            self.protocol_block,
            height=6,
            font=("Segoe UI", 10),
            wrap="word",
            relief="solid",
            bd=1
        )
        self.txt_protocol.pack(fill="both", expand=True)

        self.right_col = tk.Frame(self.row2, bg="#f5f5f5")
        self.right_col.pack(side="right", fill="both", expand=True, padx=(5, 0))

        right_col = self.right_col

        # ===== ПОРУЧЕНИЯ / ВЫБРАННЫЕ УЧАСТНИКИ СОВЕЩАНИЯ =====
        self.tasks_label = ttk.Label(right_col, text="Поручения:", style="Header.TLabel")
        self.tasks_label.pack(anchor="w", padx=5)

        self.tasks_block = tk.Frame(right_col, bg="#f5f5f5", height=90)
        self.tasks_block.pack(fill="x", expand=False, pady=(2, 0))
        self.tasks_block.pack_propagate(False)

        tasks_outer = tk.Frame(self.tasks_block, bg="#ffffff", relief="solid", bd=1)
        tasks_outer.pack(fill="both", expand=True)

        self.tasks_canvas = tk.Canvas(
            tasks_outer,
            bg="white",
            highlightthickness=0,
            height=80
        )

        tasks_scroll = ttk.Scrollbar(
            tasks_outer,
            orient="vertical",
            command=self.tasks_canvas.yview
        )

        self.tasks_container = tk.Frame(self.tasks_canvas, bg="white")

        self.tasks_container.bind(
            "<Configure>",
            lambda e:self.tasks_canvas.configure(scrollregion=self.tasks_canvas.bbox("all"))
        )

        self.tasks_canvas_window = self.tasks_canvas.create_window(
            (0, 0),
            window=self.tasks_container,
            anchor="nw"
        )

        def resize_tasks_container(event):
            self.tasks_canvas.itemconfig(self.tasks_canvas_window, width=event.width)

        self.tasks_canvas.bind("<Configure>", resize_tasks_container)
        self.tasks_canvas.configure(yscrollcommand=tasks_scroll.set)

        self.tasks_canvas.pack(side="left", fill="both", expand=True)
        # Вертикальную серую полосу прокрутки скрываем, но прокрутка колесиком мыши остается.
        # tasks_scroll intentionally not packed.

        self.tasks_rows = []
        self.agenda_speaker_rows = []
        self.tasks_mode = "tasks"

        def format_task_date(event=None):
            widget = event.widget
            digits = "".join(filter(str.isdigit, widget.get()))

            if len(digits) > 8:
                digits = digits[:8]

            if len(digits) > 4:
                formatted = f"{digits[:2]}.{digits[2:4]}.{digits[4:]}"
            elif len(digits) > 2:
                formatted = f"{digits[:2]}.{digits[2:]}"
            else:
                formatted = digits

            if widget.get() != formatted:
                widget.delete(0, "end")
                widget.insert(0, formatted)

        def _bind_tasks_mousewheel(widget):
            def _on_task_mousewheel(event):
                if getattr(event, "delta", 0):
                    self.tasks_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
                elif getattr(event, "num", None) == 4:
                    self.tasks_canvas.yview_scroll(-1, "units")
                elif getattr(event, "num", None) == 5:
                    self.tasks_canvas.yview_scroll(1, "units")
                return "break"

            widget.bind("<MouseWheel>", _on_task_mousewheel)
            widget.bind("<Button-4>", _on_task_mousewheel)
            widget.bind("<Button-5>", _on_task_mousewheel)

            for child in widget.winfo_children():
                _bind_tasks_mousewheel(child)

        def clear_tasks_container():
            for widget in self.tasks_container.winfo_children():
                widget.destroy()

        def render_task_rows():
            self.tasks_mode = "tasks"
            clear_tasks_container()

            if not self.tasks_rows:
                self.tasks_rows.append({
                    "committee":tk.StringVar(value=""),
                    "extra_committees":[],
                    "task":tk.StringVar(value=""),
                    "date":tk.StringVar(value=""),
                    "done":tk.BooleanVar(value=False)
                })

            for row_data in self.tasks_rows:
                if "committee" not in row_data:
                    row_data["committee"] = tk.StringVar(value="")
                if "extra_committees" not in row_data:
                    extras = []
                    old_second = row_data.get("second_committee")
                    try:
                        old_second_value = old_second.get().strip() if old_second else ""
                    except Exception:
                        old_second_value = ""
                    if old_second_value:
                        extras.append(tk.StringVar(value=old_second_value))
                    row_data["extra_committees"] = extras
                if "task" not in row_data:
                    row_data["task"] = tk.StringVar(value="")
                if "date" not in row_data:
                    row_data["date"] = tk.StringVar(value="")
                if "done" not in row_data:
                    row_data["done"] = tk.BooleanVar(value=False)

            protocol_mode = self.attachment_type.get() == "protocol"
            committee_values = self._get_committee_choices_for_protocol()

            if protocol_mode:
                header = tk.Frame(self.tasks_container, bg="white")
                header.pack(fill="x", pady=(0, 2), padx=5)
                header.grid_columnconfigure(3, weight=1)

                tk.Label(header, text="", width=3, bg="white").grid(row=0, column=0, padx=(0, 5), sticky="w")
                tk.Label(header, text="", width=3, bg="white").grid(row=0, column=1, padx=(0, 2), sticky="w")
                tk.Label(header, text="Комитет", bg="white", font=("Segoe UI", 9, "bold")).grid(row=0, column=2, padx=5, sticky="w")
                tk.Label(header, text="Поручение", bg="white", font=("Segoe UI", 9, "bold")).grid(row=0, column=3, padx=5, sticky="w")
                tk.Label(header, text="Срок", bg="white", font=("Segoe UI", 9, "bold")).grid(row=0, column=4, padx=5, sticky="w")
                tk.Label(header, text="", width=4, bg="white").grid(row=0, column=5, padx=5, sticky="w")
                tk.Label(header, text="", width=2, bg="white").grid(row=0, column=6, padx=5, sticky="w")

            def enable_widget_context_menu(widget, var=None):
                def do_paste(event=None):
                    try:
                        text = widget.clipboard_get()
                    except Exception:
                        return "break"
                    try:
                        widget.delete("sel.first", "sel.last")
                    except Exception:
                        pass
                    try:
                        widget.insert(tk.INSERT, text)
                    except Exception:
                        if var is not None:
                            var.set(text)
                    return "break"

                def do_copy(event=None):
                    try:
                        widget.event_generate("<<Copy>>")
                    except Exception:
                        pass
                    return "break"

                def do_cut(event=None):
                    try:
                        widget.event_generate("<<Cut>>")
                    except Exception:
                        pass
                    return "break"

                def do_select_all(event=None):
                    try:
                        widget.select_range(0, "end")
                        widget.icursor("end")
                    except Exception:
                        try:
                            widget.event_generate("<<SelectAll>>")
                        except Exception:
                            pass
                    return "break"

                def do_clear(event=None):
                    if var is not None:
                        var.set("")
                    else:
                        try:
                            widget.delete(0, "end")
                        except Exception:
                            pass
                    return "break"

                def show_menu(event):
                    menu = tk.Menu(widget, tearoff=0)
                    menu.add_command(label="Вырезать", command=do_cut)
                    menu.add_command(label="Копировать", command=do_copy)
                    menu.add_command(label="Вставить", command=do_paste)
                    menu.add_command(label="Выделить всё", command=do_select_all)
                    menu.add_separator()
                    menu.add_command(label="Очистить", command=do_clear)
                    menu.tk_popup(event.x_root, event.y_root)
                    return "break"

                widget.bind("<Control-v>", do_paste)
                widget.bind("<Control-V>", do_paste)
                widget.bind("<Shift-Insert>", do_paste)
                widget.bind("<Control-a>", do_select_all)
                widget.bind("<Control-A>", do_select_all)
                widget.bind("<Button-3>", show_menu)

            for index, row_data in enumerate(self.tasks_rows, start=1):
                row = tk.Frame(self.tasks_container, bg="white")
                row.pack(fill="x", pady=2, padx=5)

                if protocol_mode:
                    row.grid_columnconfigure(3, weight=1)
                else:
                    row.grid_columnconfigure(1, weight=1)

                tk.Label(
                    row,
                    text=f"{index}.",
                    width=3,
                    bg="white"
                ).grid(row=0, column=0, padx=(0, 5), sticky="w")

                if protocol_mode:
                    def add_extra_committee(data=row_data):
                        data.setdefault("extra_committees", [])
                        data["extra_committees"].append(tk.StringVar(value=""))
                        render_task_rows()

                    add_committee_btn = tk.Button(
                        row,
                        text="+",
                        width=2,
                        height=1,
                        font=("Segoe UI", 9, "bold"),
                        bg="white",
                        relief="solid",
                        bd=1,
                        command=add_extra_committee
                    )
                    add_committee_btn.grid(row=0, column=1, padx=(0, 2), sticky="w")

                    committee_frame = tk.Frame(row, bg="white")
                    committee_frame.grid(row=0, column=2, padx=5, sticky="ew")

                    def make_committee_line(parent, committee_var, can_remove=False, remove_index=None):
                        line = tk.Frame(parent, bg="white")
                        line.pack(fill="x", pady=(0, 2) if can_remove else 0)

                        committee_combo = ttk.Combobox(
                            line,
                            textvariable=committee_var,
                            values=committee_values,
                            state="readonly",
                            font=("Segoe UI", 10),
                            width=28
                        )
                        committee_combo.pack(side="left", fill="x", expand=True)
                        enable_widget_context_menu(committee_combo, committee_var)

                        clear_btn = tk.Button(
                            line,
                            text="×",
                            width=2,
                            height=1,
                            font=("Segoe UI", 9, "bold"),
                            bg="white",
                            relief="solid",
                            bd=1,
                            command=lambda v=committee_var: v.set("")
                        )
                        clear_btn.pack(side="left", padx=(2, 0))

                        if can_remove:
                            def remove_extra(data=row_data, idx=remove_index):
                                try:
                                    data["extra_committees"].pop(idx)
                                except Exception:
                                    pass
                                render_task_rows()

                            remove_btn = tk.Button(
                                line,
                                text="−",
                                width=2,
                                height=1,
                                font=("Segoe UI", 9, "bold"),
                                bg="white",
                                relief="solid",
                                bd=1,
                                command=remove_extra
                            )
                            remove_btn.pack(side="left", padx=(2, 0))

                    make_committee_line(committee_frame, row_data["committee"])

                    for extra_index, extra_var in enumerate(list(row_data.get("extra_committees", []))):
                        make_committee_line(
                            committee_frame,
                            extra_var,
                            can_remove=True,
                            remove_index=extra_index
                        )

                    task_entry = tk.Entry(
                        row,
                        textvariable=row_data["task"],
                        font=("Segoe UI", 10),
                        width=48
                    )
                    task_entry.grid(row=0, column=3, padx=5, sticky="ew")
                    enable_widget_context_menu(task_entry, row_data["task"])

                    date_entry = tk.Entry(
                        row,
                        textvariable=row_data["date"],
                        font=("Segoe UI", 10),
                        width=12
                    )
                    date_entry.grid(row=0, column=4, padx=5, sticky="w")
                    enable_widget_context_menu(date_entry, row_data["date"])

                    status_label = tk.Label(
                        row,
                        text="",
                        width=2,
                        bg="white",
                        font=("Segoe UI", 12, "bold")
                    )

                    def update_status_label(label=status_label, data=row_data):
                        date_text = data["date"].get().strip()
                        is_done = data["done"].get()

                        if date_text and is_done:
                            label.config(text="✓", fg="#148a2b")
                        elif date_text and not is_done:
                            label.config(text="!", fg="#c00000")
                        else:
                            label.config(text="", fg="#333333")

                    def on_date_change(event=None, label=status_label, data=row_data):
                        format_task_date(event)
                        update_status_label(label, data)

                    date_entry.bind("<KeyRelease>", on_date_change)

                    chk = tk.Checkbutton(
                        row,
                        variable=row_data["done"],
                        bg="white",
                        command=update_status_label
                    )
                    chk.grid(row=0, column=5, padx=5, sticky="w")

                    status_label.grid(row=0, column=6, padx=(0, 5), sticky="w")
                    update_status_label()
                else:
                    task_entry = tk.Entry(
                        row,
                        textvariable=row_data["task"],
                        font=("Segoe UI", 10),
                        width=70
                    )
                    task_entry.grid(row=0, column=1, padx=5, sticky="ew")
                    enable_widget_context_menu(task_entry, row_data["task"])

                    date_entry = tk.Entry(
                        row,
                        textvariable=row_data["date"],
                        font=("Segoe UI", 10),
                        width=12
                    )
                    date_entry.grid(row=0, column=2, padx=5, sticky="w")
                    date_entry.bind("<KeyRelease>", format_task_date)
                    enable_widget_context_menu(date_entry, row_data["date"])

                    chk = tk.Checkbutton(
                        row,
                        variable=row_data["done"],
                        bg="white"
                    )
                    chk.grid(row=0, column=3, padx=5, sticky="w")

            self.tasks_container.update_idletasks()
            self.tasks_canvas.configure(scrollregion=self.tasks_canvas.bbox("all"))
            _bind_tasks_mousewheel(self.tasks_container)

        def add_task_row(task_text="", date_text="", done=False, committee_text="", second_committee_text=""):
            extra_values = []
            if isinstance(second_committee_text, (list, tuple)):
                extra_values = [str(value or "") for value in second_committee_text if str(value or "").strip()]
            else:
                second_committee_text = str(second_committee_text or "")
                if "§§" in second_committee_text:
                    extra_values = [value.strip() for value in second_committee_text.split("§§") if value.strip()]
                elif second_committee_text.strip():
                    extra_values = [second_committee_text.strip()]

            self.tasks_rows.append({
                "committee":tk.StringVar(value=committee_text),
                "extra_committees":[tk.StringVar(value=value) for value in extra_values],
                "task":tk.StringVar(value=task_text),
                "date":tk.StringVar(value=date_text),
                "done":tk.BooleanVar(value=done)
            })

            if self.tasks_mode == "tasks":
                render_task_rows()
                self.tasks_canvas.yview_moveto(1.0)

        def remove_task_row():
            if self.tasks_mode == "tasks" and self.tasks_rows:
                self.tasks_rows.pop()
                render_task_rows()

        def render_agenda_speaker_rows():
            self.tasks_mode = "agenda"
            clear_tasks_container()

            for index, row_data in enumerate(self.agenda_speaker_rows, start=1):
                row = tk.Frame(self.tasks_container, bg="white")
                row.pack(fill="x", pady=2, padx=5)

                row.grid_columnconfigure(1, weight=1)
                row.grid_columnconfigure(2, weight=1)

                tk.Label(
                    row,
                    text=f"{index}.",
                    width=3,
                    bg="white"
                ).grid(row=0, column=0, padx=(0, 5), sticky="w")

                fio_entry = tk.Entry(
                    row,
                    textvariable=row_data["fio"],
                    font=("Segoe UI", 10)
                )
                fio_entry.grid(row=0, column=1, padx=5, sticky="ew")

                position_entry = tk.Entry(
                    row,
                    textvariable=row_data["position"],
                    font=("Segoe UI", 10)
                )
                position_entry.grid(row=0, column=2, padx=5, sticky="ew")

                chk = tk.Checkbutton(
                    row,
                    variable=row_data["checked"],
                    bg="white"
                )
                chk.grid(row=0, column=3, padx=5, sticky="w")

                question_entry = tk.Entry(
                    row,
                    textvariable=row_data["question_number"],
                    font=("Segoe UI", 10),
                    width=4,
                    justify="center"
                )
                question_entry.grid(row=0, column=4, padx=(2, 5), sticky="w")

            self.tasks_container.update_idletasks()
            self.tasks_canvas.configure(scrollregion=self.tasks_canvas.bbox("all"))

        self.render_task_rows = render_task_rows
        self.render_agenda_speaker_rows = render_agenda_speaker_rows
        self.add_task_row = add_task_row
        self.remove_task_row = remove_task_row

        self.task_btns = tk.Frame(right_col, bg="#f5f5f5")
        self.task_btns.pack(fill="x")

        tk.Button(
            self.task_btns,
            text="+",
            command=self.add_task_row
        ).pack(side="left", padx=5, pady=3)

        tk.Button(
            self.task_btns,
            text="-",
            command=self.remove_task_row
        ).pack(side="left", padx=5, pady=3)

        add_task_row()

        # ===== ИТОГОВОЕ ПОРУЧЕНИЕ ПО ИНФОРМИРОВАНИЮ =====
        self.protocol_report_enabled = tk.BooleanVar(value=False)
        self.protocol_report_text = tk.StringVar(
            value="О проделанной работе проинформировать правовой комитет администрации города Барнаула до "
        )
        self.protocol_report_date = tk.StringVar(value="")

        self.protocol_report_frame = tk.Frame(right_col, bg="#f5f5f5")
        self.protocol_report_frame.pack(fill="x", pady=(2, 0))
        self.protocol_report_frame.grid_columnconfigure(0, weight=1)

        self.protocol_report_entry = tk.Entry(
            self.protocol_report_frame,
            textvariable=self.protocol_report_text,
            font=("Segoe UI", 10),
            relief="solid",
            bd=1,
            state="disabled",
            disabledbackground="#e8e8e8",
            disabledforeground="#666666"
        )
        self.protocol_report_entry.grid(row=0, column=0, padx=(0, 5), sticky="ew")

        self.protocol_report_date_entry = tk.Entry(
            self.protocol_report_frame,
            textvariable=self.protocol_report_date,
            font=("Segoe UI", 10),
            width=12,
            relief="solid",
            bd=1
        )
        self.protocol_report_date_entry.grid(row=0, column=1, padx=5, sticky="w")
        self.protocol_report_date_entry.bind("<KeyRelease>", format_task_date)

        def update_protocol_report_state():
            if self.protocol_report_enabled.get():
                self.protocol_report_entry.config(
                    state="normal",
                    bg="white",
                    fg="black"
                )
            else:
                self.protocol_report_entry.config(
                    state="disabled",
                    disabledbackground="#e8e8e8",
                    disabledforeground="#666666"
                )

        self.update_protocol_report_state = update_protocol_report_state

        self.protocol_report_check = tk.Checkbutton(
            self.protocol_report_frame,
            variable=self.protocol_report_enabled,
            bg="#f5f5f5",
            command=update_protocol_report_state
        )
        self.protocol_report_check.grid(row=0, column=2, padx=(5, 0), sticky="w")

        try:
            self.protocol_report_entry.bind("<Control-a>", lambda e: (self.protocol_report_entry.select_range(0, "end"), "break")[-1])
            self.protocol_report_entry.bind("<Control-A>", lambda e: (self.protocol_report_entry.select_range(0, "end"), "break")[-1])
            self.protocol_report_date_entry.bind("<Control-a>", lambda e: (self.protocol_report_date_entry.select_range(0, "end"), "break")[-1])
            self.protocol_report_date_entry.bind("<Control-A>", lambda e: (self.protocol_report_date_entry.select_range(0, "end"), "break")[-1])
        except Exception:
            pass

        def enable_report_context_menu(entry_widget):
            def paste_text(event=None):
                try:
                    text = entry_widget.clipboard_get()
                    try:
                        entry_widget.delete("sel.first", "sel.last")
                    except Exception:
                        pass
                    entry_widget.insert(tk.INSERT, text)
                except Exception:
                    pass
                return "break"

            def select_all(event=None):
                try:
                    entry_widget.select_range(0, "end")
                    entry_widget.icursor("end")
                except Exception:
                    pass
                return "break"

            def show_menu(event):
                menu = tk.Menu(entry_widget, tearoff=0)
                menu.add_command(label="Вырезать", command=lambda: entry_widget.event_generate("<<Cut>>"))
                menu.add_command(label="Копировать", command=lambda: entry_widget.event_generate("<<Copy>>"))
                menu.add_command(label="Вставить", command=paste_text)
                menu.add_command(label="Выделить всё", command=select_all)
                menu.add_separator()
                menu.add_command(label="Очистить", command=lambda: entry_widget.delete(0, "end"))
                menu.tk_popup(event.x_root, event.y_root)
                return "break"

            entry_widget.bind("<Control-v>", paste_text)
            entry_widget.bind("<Control-V>", paste_text)
            entry_widget.bind("<Shift-Insert>", paste_text)
            entry_widget.bind("<Control-a>", select_all)
            entry_widget.bind("<Control-A>", select_all)
            entry_widget.bind("<Button-3>", show_menu)

        enable_report_context_menu(self.protocol_report_entry)
        enable_report_context_menu(self.protocol_report_date_entry)

        update_protocol_report_state()

        # ===== ПАНЕЛЬ УПРАВЛЕНИЯ ДОКУМЕНТАМИ =====
        # Кнопка "Сформировать документ" слева, документы и кнопки
        # "Прикрепить" / "Открыть" — справа, на одном уровне.
        self.document_bar = tk.Frame(form_container, bg="#f5f5f5")
        self.document_bar.pack(fill="x", pady=(4, 4))

        self.btn_generate_doc = ttk.Button(
            self.document_bar,
            text="📄 Сформировать документ",
            command=self.generate_document,
            style="Gray.TButton"
        )
        self.btn_generate_doc.pack(side="left", padx=5)

        self.file_controls_inner = tk.Frame(self.document_bar, bg="#f5f5f5")
        self.file_controls_inner.pack(side="right", padx=5)

        ttk.Label(
            self.file_controls_inner,
            text="Документы:",
            style="Header.TLabel"
        ).pack(side="left", padx=(5, 6))

        self.lbl_file = tk.Label(
            self.file_controls_inner,
            textvariable=self.attachment_var,
            bg="white",
            fg="#555",
            anchor="w",
            wraplength=240,
            justify="left",
            relief="solid",
            bd=1,
            width=30,
            height=1
        )
        self.lbl_file.pack(side="left", padx=(0, 6))

        self.btn_attach_file = ttk.Button(
            self.file_controls_inner,
            text="📎 Прикрепить",
            command=self.attach_file,
            style="Modern.TButton"
        )
        self.btn_attach_file.pack(side="left", padx=(0, 5))

        self.btn_open_file = ttk.Button(
            self.file_controls_inner,
            text="📂 Открыть",
            command=self.open_attached_file,
            style="Modern.TButton"
        )
        self.btn_open_file.pack(side="left", padx=(0, 5))

        self.lbl_file.bind("<Double-1>", lambda event: self.open_attached_file())

        # Кнопки выбора участников.
        self.participants_select_frame = tk.Frame(form_container, bg="#f5f5f5")
        self.participants_select_frame.pack(fill="x", pady=(0, 2))

        self.btn_select_local = ttk.Button(
            self.participants_select_frame,
            text="👥 Органы МСУ",
            command=self.open_user_selector,
            style="Modern.TButton"
        )
        self.btn_select_local.pack(side="left", padx=5, pady=2)

        self.btn_select_invited = ttk.Button(
            self.participants_select_frame,
            text="👥 Приглашенные",
            command=self.open_invited_selector,
            style="Modern.TButton"
        )
        self.btn_select_invited.pack(side="left", padx=5, pady=2)

        # Кнопки выбора типа документа уже созданы в верхней строке,
        # рядом с текстом "Выберите вид создаваемого документа:".

        def update_doc_type_buttons():
            selected = self.attachment_type.get()

            if selected != "agenda" and hasattr(self, "txt_agenda"):
                if self.agenda_placeholder_active:
                    self.hide_agenda_placeholder()

            for doc_type, btn in self.doc_type_buttons.items():
                if doc_type == selected:
                    btn.configure(style="Selected.TButton")
                else:
                    btn.configure(style="Modern.TButton")

        def set_doc_type(doc_type):
            self.attachment_type.set(doc_type)
            update_doc_type_buttons()
            self.update_visible_fields_by_doc_type()

        btn_participants = ttk.Button(
            self.doc_buttons_frame,
            text="Список участников",
            style="Modern.TButton",
            command=lambda:set_doc_type("participants")
        )
        btn_participants.pack(side="left", padx=5, pady=2)
        self.doc_type_buttons["participants"] = btn_participants

        btn_agenda = ttk.Button(
            self.doc_buttons_frame,
            text="Повестка",
            style="Modern.TButton",
            command=lambda:set_doc_type("agenda")
        )
        btn_agenda.pack(side="left", padx=5, pady=2)
        self.doc_type_buttons["agenda"] = btn_agenda

        btn_telegram = ttk.Button(
            self.doc_buttons_frame,
            text="Телефонограмма",
            style="Modern.TButton",
            command=lambda:set_doc_type("telegram")
        )
        btn_telegram.pack(side="left", padx=5, pady=2)
        self.doc_type_buttons["telegram"] = btn_telegram

        btn_protocol = ttk.Button(
            self.doc_buttons_frame,
            text="Протокол",
            style="Modern.TButton",
            command=lambda:set_doc_type("protocol")
        )
        btn_protocol.pack(side="left", padx=5, pady=2)
        self.doc_type_buttons["protocol"] = btn_protocol

        # ===== Почта для докладов =====
        self.email_block = tk.Frame(form_container, bg="#f5f5f5")
        self.email_block.pack(fill="x", pady=2)

        ttk.Label(self.email_block, text="Почта для докладов:", width=20, anchor="e").pack(side="left", padx=5)
        self.ent_email = tk.Entry(self.email_block, font=("Segoe UI", 10), width=40)
        self.ent_email.insert(0, "fedorova-en@barnaul-adm.ru")
        self.ent_email.pack(side="left", padx=5, fill="x", expand=True)

        # ===== Блок реквизитов для телефонограммы =====
        self.transfer_block = tk.Frame(form_container, bg="#f5f5f5")
        self.transfer_block.pack(fill="x", pady=2)

        ttk.Label(self.transfer_block, text="Электронная почта:", width=16, anchor="e").pack(side="left", padx=5)
        tk.Entry(self.transfer_block, textvariable=self.transfer_email, font=("Segoe UI", 10), width=28).pack(side="left", padx=5)

        ttk.Label(self.transfer_block, text="Передала:", width=10, anchor="e").pack(side="left", padx=(10, 3))
        tk.Entry(self.transfer_block, textvariable=self.transfer_fio, font=("Segoe UI", 10), width=24).pack(side="left", padx=5)

        ttk.Label(self.transfer_block, text="Тел.:", anchor="e").pack(side="left", padx=(10, 3))
        tk.Entry(self.transfer_block, textvariable=self.transfer_phone, font=("Segoe UI", 10), width=14).pack(side="left", padx=5)

        ttk.Label(self.transfer_block, text="ФИО:", anchor="e").pack(side="left", padx=(10, 3))
        tk.Entry(self.transfer_block, textvariable=self.telegram_sign_fio, font=("Segoe UI", 10), width=16).pack(side="left", padx=5)

        # Совместимость со старым кодом: эти блоки больше не используются
        # для размещения, но методы могут проверять их наличие.
        self.telegram_number_frame = self.top_telegram_number_block
        self.transfer_email_block = self.transfer_block

        # ===== Поля подписи для повестки =====
        self.agenda_sign_block = tk.Frame(form_container, bg="#f5f5f5")
        self.agenda_sign_block.pack(fill="x", pady=(4, 0))

        ttk.Label(self.agenda_sign_block, text="Должность:", width=14, anchor="e").pack(side="left", padx=5)
        self.ent_agenda_sign_position = tk.Entry(
            self.agenda_sign_block,
            textvariable=self.agenda_sign_position,
            font=("Segoe UI", 10),
            width=45
        )
        self.ent_agenda_sign_position.pack(side="left", padx=5, fill="x", expand=True)

        ttk.Label(self.agenda_sign_block, text="ФИО:", width=8, anchor="e").pack(side="left", padx=5)
        self.ent_agenda_sign_fio = tk.Entry(
            self.agenda_sign_block,
            textvariable=self.agenda_sign_fio,
            font=("Segoe UI", 10),
            width=22
        )
        self.ent_agenda_sign_fio.pack(side="left", padx=5)

        # ===== Поле "Протокол ведет" для списка участников =====
        self.protocol_keeper_block = tk.Frame(form_container, bg="#f5f5f5")
        self.protocol_keeper_block.pack(fill="x", pady=(4, 0))

        ttk.Label(
            self.protocol_keeper_block,
            text="Протокол ведет:",
            width=20,
            anchor="e"
        ).pack(side="left", padx=5)

        self.ent_protocol_keeper = tk.Entry(
            self.protocol_keeper_block,
            textvariable=self.protocol_keeper,
            font=("Segoe UI", 10),
            width=45
        )
        self.ent_protocol_keeper.pack(side="left", padx=5, fill="x", expand=True)

        # ===== Поля протокола в основном окне, а не в отдельном окне =====
        self.protocol_inline_block = tk.Frame(form_container, bg="#f5f5f5")
        self.protocol_inline_block.pack(fill="x", pady=(4, 0))

        ttk.Label(self.protocol_inline_block, text="Председательствующий:", width=21, anchor="e").pack(side="left", padx=5)
        tk.Entry(self.protocol_inline_block, textvariable=self.protocol_chair_fio, font=("Segoe UI", 10), width=24).pack(side="left", padx=5)

        ttk.Label(self.protocol_inline_block, text="Должность:", width=11, anchor="e").pack(side="left", padx=5)
        tk.Entry(self.protocol_inline_block, textvariable=self.protocol_chair_position, font=("Segoe UI", 10), width=58).pack(side="left", padx=5, fill="x", expand=True)


        update_doc_type_buttons()

        self.list_container = tk.Frame(self, bg="#f5f5f5")
        self.list_container.pack(fill="both", expand=True, padx=10, pady=(2, 6))
        list_container = self.list_container

        ttk.Label(list_container, text="Архив совещаний", style="Header.TLabel").pack(anchor="w", pady=(0, 5))

        columns = ("id", "date_val", "time_val", "title", "participants_preview")
        self.tree = ttk.Treeview(list_container, columns=columns, show='headings', selectmode="browse")

        self.tree.heading("id", text="№ п/п")
        self.tree.heading("date_val", text="Дата")
        self.tree.heading("time_val", text="Время")
        self.tree.heading("title", text="Тема совещания")
        self.tree.heading("participants_preview", text="Участники")

        self.tree.column("id", width=80, anchor="center")
        self.tree.column("date_val", width=80, anchor="center")
        self.tree.column("time_val", width=60, anchor="center")
        self.tree.column("title", width=400, anchor="w")
        self.tree.column("participants_preview", width=150, anchor="w")

        scroll_y = ttk.Scrollbar(list_container, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=scroll_y.set)

        self.tree.pack(side="left", fill="both", expand=True)
        scroll_y.pack(side="right", fill="y")

        # ===== Кнопки управления под таблицей =====
        self.bottom_action_bar = tk.Frame(self, bg="#f5f5f5")
        self.bottom_action_bar.pack(fill="x", padx=10, pady=(0, 10))

        ttk.Button(
            self.bottom_action_bar,
            text="💾 Сохранить",
            command=self.save_meeting,
            style="Modern.TButton"
        ).pack(side="left", padx=2)

        ttk.Button(
            self.bottom_action_bar,
            text="🗑️ Удалить",
            command=self.delete_meeting,
            style="Modern.TButton"
        ).pack(side="left", padx=2)

        ttk.Button(
            self.bottom_action_bar,
            text="❌ Закрыть",
            command=self.destroy,
            style="Modern.TButton"
        ).pack(side="right", padx=2)

        self.tree.bind("<Double-1>", self.on_select)
        self.load_meetings()
        self.update_visible_fields_by_doc_type()

        # Если тип документа был выбран в стартовом мини-окне, сразу открываем
        # нужный раздел. Иначе оставляем старый резервный вариант.
        if initial_doc_type:
            self.after(10, lambda value=initial_doc_type: self._apply_document_type_from_start(value))
        else:
            self.after(50, self.show_initial_doc_type_choice)


    def _apply_document_type_from_start(self, doc_type):
        """
        Применяет выбранный в стартовом мини-окне тип документа.
        Работает так же, как нажатие верхних кнопок в основном окне.
        """
        self.attachment_type.set(doc_type)

        selected = self.attachment_type.get()

        for button_type, button in getattr(self, "doc_type_buttons", {}).items():
            try:
                if button_type == selected:
                    button.configure(style="Selected.TButton")
                else:
                    button.configure(style="Modern.TButton")
            except Exception:
                pass

        self.update_visible_fields_by_doc_type()
        self.deiconify()
        self.lift()
        self.focus_force()

    def show_initial_doc_type_choice(self):
        """
        Стартовое мини-окно выбора типа документа.
        Не дает пользователю увидеть нейтральное состояние основного окна.
        """
        if getattr(self, "_initial_doc_type_chosen", False):
            return

        self.withdraw()

        dialog_parent = getattr(self, "parent_main", None) or self.master or self
        dialog = tk.Toplevel(dialog_parent)
        dialog.title("Выбор документа")
        dialog.configure(bg="#f4f6f9")
        dialog.resizable(False, False)
        try:
            dialog.transient(dialog_parent)
        except Exception:
            pass
        dialog.grab_set()

        try:
            apply_modern_style(dialog)
        except Exception:
            pass

        container = tk.Frame(dialog, bg="#f4f6f9")
        container.pack(fill="both", expand=True, padx=22, pady=18)

        tk.Label(
            container,
            text="Выберите тип создаваемого документа",
            font=("Segoe UI", 13, "bold"),
            bg="#f4f6f9",
            fg="#111111"
        ).pack(anchor="center", pady=(0, 16))

        buttons_frame = tk.Frame(container, bg="#f4f6f9")
        buttons_frame.pack(fill="both", expand=True)

        doc_buttons = [
            ("Список участников", "participants"),
            ("Повестка", "agenda"),
            ("Телефонограмма", "telegram"),
            ("Протокол", "protocol"),
        ]

        def choose(doc_type):
            self._initial_doc_type_chosen = True
            try:
                dialog.grab_release()
            except Exception:
                pass
            dialog.destroy()
            self._apply_document_type_from_start(doc_type)

        for title, doc_type in doc_buttons:
            ttk.Button(
                buttons_frame,
                text=title,
                style="Modern.TButton",
                command=lambda value=doc_type: choose(value)
            ).pack(fill="x", pady=5, ipady=4)

        def close_without_choice():
            try:
                dialog.grab_release()
            except Exception:
                pass
            dialog.destroy()
            self.destroy()

        dialog.protocol("WM_DELETE_WINDOW", close_without_choice)
        dialog.bind("<Escape>", lambda event: close_without_choice())

        dialog.update_idletasks()
        width = max(dialog.winfo_width(), 430)
        height = max(dialog.winfo_height(), 260)
        screen_w = dialog.winfo_screenwidth()
        screen_h = dialog.winfo_screenheight()
        x = (screen_w - width) // 2
        y = (screen_h - height) // 2
        dialog.geometry(f"{width}x{height}+{x}+{y}")

        dialog.lift()
        dialog.focus_force()


    def _normalize_text(self, text):
        text = str(text or "").strip().lower().replace("ё", "е")
        return " ".join(text.split())

    def _short_fio_initials_for_label(self, full_name):
        """
        Белоцерковец Максим Александрович -> Белоцерковец М.А.
        """
        parts = str(full_name or "").strip().split()

        if len(parts) >= 3:
            return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."

        return str(full_name or "").strip()

    def _extract_committee_direction_from_position(self, position):
        """
        Из должности председателя комитета берет направление после слова "комитета".

        Пример:
        "Председатель комитета по общественным связям и безопасности"
        -> "по общественным связям и безопасности"
        """
        raw = str(position or "").strip()
        norm = self._normalize_text(raw)

        if "председатель" not in norm or "комитет" not in norm:
            return ""

        marker_variants = ["комитета", "комитет"]

        for marker in marker_variants:
            idx = norm.find(marker)

            if idx == -1:
                continue

            direction = raw[idx + len(marker):].strip(" ,.-")

            if direction:
                if not self._normalize_text(direction).startswith("по "):
                    direction = "по " + direction

                return direction

            return "по общим вопросам"

        return ""

    def _get_committee_choices_for_protocol(self):
        """
        Для поля "Комитет" в блоке поручений протокола:
        берет все комитеты из должностей сотрудников списка "Органы МСУ"
        и добавляет в скобках председателя комитета в формате Фамилия И.О.
        """
        try:
            people = self._load_local_people_records()

            if not people:
                people = get_all_employees()
        except Exception:
            people = []

        values = []
        seen = set()

        for person in people:
            position = (
                person.get("position")
                or person.get("dolzhnost")
                or person.get("job_title")
                or ""
            )
            full_name = (
                person.get("full_name")
                or person.get("fio")
                or person.get("name")
                or ""
            )

            direction = self._extract_committee_direction_from_position(position)

            if not direction:
                continue

            label = direction

            short_fio = self._short_fio_initials_for_label(full_name)

            if short_fio:
                label = f"{direction} ({short_fio})"

            key = self._normalize_text(label)

            if key in seen:
                continue

            seen.add(key)
            values.append(label)

        return sorted(values, key=self._normalize_text)

    def _get_task_committees_from_row(self, task_row):
        """
        Возвращает все выбранные комитеты из строки поручения.
        Поддерживает старый формат second_committee и новый список extra_committees.
        """
        committees = []

        try:
            first = task_row.get("committee").get().strip() if task_row.get("committee") else ""
        except Exception:
            first = ""

        if first:
            committees.append(first)

        try:
            extras = task_row.get("extra_committees", []) or []
        except Exception:
            extras = []

        for extra in extras:
            try:
                value = extra.get().strip()
            except Exception:
                value = str(extra or "").strip()

            if value and value not in committees:
                committees.append(value)

        # Совместимость со старыми файлами, где был только second_committee.
        try:
            second = task_row.get("second_committee").get().strip() if task_row.get("second_committee") else ""
        except Exception:
            second = ""

        if second and second not in committees:
            committees.append(second)

        return committees

    def _format_protocol_committee_assignment(self, first_committee=None, second_committee=""):
        """
        Формирует начало пункта решения протокола по выбранным комитетам.

        В поле хранится: по энергоресурсам и газификации (Крюков А.В.)
        В документе: Комитету по энергоресурсам и газификации (Крюков А.В.)
        Если комитетов несколько, далее идет:
        совместно с комитетом по ... (ФИО)
        """
        if isinstance(first_committee, (list, tuple)):
            committees = [str(value or "").strip() for value in first_committee if str(value or "").strip()]
        else:
            committees = [str(first_committee or "").strip()] if str(first_committee or "").strip() else []

            for value in str(second_committee or "").split("§§"):
                value = value.strip()
                if value:
                    committees.append(value)

        unique_committees = []

        for committee in committees:
            if committee and committee not in unique_committees:
                unique_committees.append(committee)

        if not unique_committees:
            return ""

        result = f"Комитету {unique_committees[0]}"

        for committee in unique_committees[1:]:
            result += f" совместно с комитетом {committee}"

        return result

    def _load_local_people_records(self):
        """
        Загружает "Органы МСУ" из настроек:
        таблица meeting_participants, category='msu_ip'.
        """
        try:
            with get_connection() as conn:
                cursor = conn.cursor()

                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS meeting_participants (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        category TEXT NOT NULL,
                        full_name TEXT NOT NULL,
                        position TEXT DEFAULT '',
                        leadership TEXT DEFAULT '',
                        sort_order INTEGER DEFAULT 999
                    )
                    """
                )

                cursor.execute(
                    """
                    SELECT id, full_name, position, leadership, sort_order
                    FROM meeting_participants
                    WHERE category = 'msu_ip'
                    ORDER BY sort_order, full_name
                    """
                )

                rows = cursor.fetchall()
                result = []

                for row in rows:
                    result.append({
                        "id": row["id"] if hasattr(row, "keys") else row[0],
                        "full_name": row["full_name"] if hasattr(row, "keys") else row[1],
                        "position": row["position"] if hasattr(row, "keys") else row[2],
                        "leadership": row["leadership"] if hasattr(row, "keys") else row[3],
                        "sort_order": row["sort_order"] if hasattr(row, "keys") else row[4],
                        "source_table": "meeting_participants",
                    })

                return result

        except Exception:
            return []

    def _load_invited_people_records(self):
        """
        Загружает приглашенных из новой таблицы meeting_participants,
        которая заполняется в настройках: Участники совещаний -> Список приглашенных.
        Возвращает записи с полем leadership, чтобы в окне выбора и при
        формировании телефонограммы работала подчиненность.
        """
        try:
            with get_connection() as conn:
                cursor = conn.cursor()

                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS meeting_participants (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        category TEXT NOT NULL,
                        full_name TEXT NOT NULL,
                        position TEXT DEFAULT '',
                        leadership TEXT DEFAULT '',
                        sort_order INTEGER DEFAULT 999
                    )
                    """
                )

                cursor.execute(
                    """
                    SELECT id, full_name, position, leadership, sort_order
                    FROM meeting_participants
                    WHERE category = 'invited_ip'
                    ORDER BY sort_order, full_name
                    """
                )

                rows = cursor.fetchall()
                result = []

                for row in rows:
                    result.append({
                        "id": row["id"] if hasattr(row, "keys") else row[0],
                        "full_name": row["full_name"] if hasattr(row, "keys") else row[1],
                        "position": row["position"] if hasattr(row, "keys") else row[2],
                        "leadership": row["leadership"] if hasattr(row, "keys") else row[3],
                        "sort_order": row["sort_order"] if hasattr(row, "keys") else row[4],
                    })

                return result

        except Exception:
            return []

    def _ensure_meetings_extra_columns(self):
        """
        Добавляет в таблицу meetings поле protocol_keeper,
        чтобы значение 'Протокол ведёт' сохранялось и загружалось.
        """
        try:
            with get_connection() as conn:
                cursor = conn.cursor()

                cursor.execute("PRAGMA table_info(meetings)")
                existing_columns = [row[1] for row in cursor.fetchall()]

                if "protocol_keeper" not in existing_columns:
                    cursor.execute(
                        "ALTER TABLE meetings ADD COLUMN protocol_keeper TEXT"
                    )

                if "invited_participants" not in existing_columns:
                    cursor.execute(
                        "ALTER TABLE meetings ADD COLUMN invited_participants TEXT"
                    )

                conn.commit()

        except Exception as e:
            print(f"Ошибка при добавлении колонки protocol_keeper в meetings: {e}")

    def _save_meeting_extra_fields(self, meeting_id):
        """
        Принудительно сохраняет поле 'Протокол ведёт',
        даже если meetings_add / meetings_update в database.py его не обрабатывают.
        """
        if not meeting_id:
            return

        try:
            with get_connection() as conn:
                cursor = conn.cursor()

                cursor.execute("PRAGMA table_info(meetings)")
                existing_columns = [row[1] for row in cursor.fetchall()]

                if "protocol_keeper" not in existing_columns:
                    cursor.execute(
                        "ALTER TABLE meetings ADD COLUMN protocol_keeper TEXT"
                    )

                cursor.execute(
                    """
                    UPDATE meetings
                    SET protocol_keeper = ?,
                        invited_participants = ?
                    WHERE id = ?
                    """,
                    (
                        self.protocol_keeper.get().strip(),
                        "\n".join(self.temp_invited_participants),
                        meeting_id
                    )
                )

                conn.commit()

        except Exception as e:
            print(f"Ошибка сохранения protocol_keeper: {e}")

    def _get_last_saved_meeting_id(self):
        """
        Находит последнюю сохраненную запись.
        Нужно на случай, если meetings_add(data) не возвращает id.
        """
        try:
            with get_connection() as conn:
                cursor = conn.cursor()

                cursor.execute(
                    """
                    SELECT id
                    FROM meetings
                    WHERE title = ?
                      AND date_val = ?
                      AND time_val = ?
                    ORDER BY id DESC
                    LIMIT 1
                    """,
                    (
                        self.ent_title.get().strip(),
                        self.ent_date.get().strip(),
                        self.ent_time.get().strip()
                    )
                )

                row = cursor.fetchone()

                if row:
                    return row[0]

        except Exception as e:
            print(f"Ошибка получения последнего id совещания: {e}")

        return None

    def show_agenda_placeholder(self):
        text = self.txt_agenda.get("1.0", "end").strip()

        if not text:
            self.agenda_placeholder_active = True
            self.txt_agenda.config(state="normal")
            self.txt_agenda.delete("1.0", "end")
            self.txt_agenda.insert("1.0", self.agenda_placeholder_text)
            self.txt_agenda.config(fg="#9a9a9a")

    def hide_agenda_placeholder(self):
        if self.agenda_placeholder_active:
            self.txt_agenda.config(state="normal")
            self.txt_agenda.delete("1.0", "end")
            self.txt_agenda.config(fg="black")
            self.agenda_placeholder_active = False

    def on_agenda_click(self, event=None):
        if self.agenda_placeholder_active:
            self.hide_agenda_placeholder()

    def on_agenda_key(self, event=None):
        if self.agenda_placeholder_active:
            self.hide_agenda_placeholder()

    def on_agenda_focus_out(self, event=None):
        text = self.txt_agenda.get("1.0", "end").strip()

        if not text and self.attachment_type.get() == "agenda":
            self.show_agenda_placeholder()

    def get_agenda_text_for_save(self):
        if self.agenda_placeholder_active:
            return ""

        text = self.txt_agenda.get("1.0", "end").strip()

        if text == self.agenda_placeholder_text:
            return ""

        return text

    def _sync_agenda_hidden_text(self):
        """
        Синхронизирует скрытый self.txt_agenda со строками блока
        'Вопрос в повестке', чтобы старый код не ломался.
        """
        try:
            text = self.get_agenda_text_for_save()
            self.txt_agenda.delete("1.0", "end")
            self.txt_agenda.insert("1.0", text)
        except Exception:
            pass

    def _set_agenda_questions_from_text(self, text):
        """
        Загружает сохраненный текст повестки обратно в строки
        блока 'Вопрос в повестке'.
        """
        if not hasattr(self, "agenda_question_rows"):
            return

        self.agenda_question_rows.clear()

        lines = [
            line.strip()
            for line in str(text or "").split("\n")
            if line.strip()
        ]

        if not lines:
            lines = [""]

        for line in lines:
            self.agenda_question_rows.append({
                "text": tk.StringVar(value=line)
            })

        if hasattr(self, "render_agenda_questions"):
            self.render_agenda_questions()

        self._sync_agenda_hidden_text()

    def get_agenda_text_for_save(self):
        """
        Возвращает текст из блока 'Вопрос в повестке'.
        Каждая строка — отдельный вопрос.
        """
        if not hasattr(self, "agenda_question_rows"):
            return ""

        lines = []

        for row in self.agenda_question_rows:
            text = row["text"].get().strip()

            if text:
                lines.append(text)

        return "\n".join(lines)

    def get_agenda_questions_for_doc(self):
        """
        Возвращает список вопросов для документа 'Повестка'.
        """
        text = self.get_agenda_text_for_save()

        questions = [
            line.strip()
            for line in text.split("\n")
            if line.strip()
        ]

        if not questions:
            questions = [""]

        return questions

    def _normalize_agenda_question_number(self, value):
        """
        Возвращает номер вопроса повестки как int.
        Если номер не указан, считаем, что участник относится к 1-му вопросу.
        """
        value = str(value or "").strip()

        if not value:
            return 1

        digits = "".join(ch for ch in value if ch.isdigit())

        if not digits:
            return 1

        try:
            return int(digits)
        except Exception:
            return 1

    def _get_checked_agenda_people_by_question_for_doc(self):
        """
        Группирует отмеченных участников блока
        'Выбранные участники совещания' по номеру вопроса повестки.

        source='local'    -> Докладывает/Докладывают
        source='invited'  -> Информирует/Информируют
        """
        result = {}

        for row in getattr(self, "agenda_speaker_rows", []):
            if not row["checked"].get():
                continue

            fio = row["fio"].get().strip()
            position = row["position"].get().strip()

            if not fio:
                continue

            question_number = self._normalize_agenda_question_number(
                row.get("question_number", tk.StringVar(value="1")).get()
            )

            source = row.get("source", "local")

            result.setdefault(question_number, {
                "local": [],
                "invited": []
            })

            target_key = "invited" if source == "invited" else "local"

            result[question_number][target_key].append({
                "fio": fio,
                "dolzhnost": position or "Не указана"
            })

        for question_number, groups in result.items():
            groups["local"] = self._sort_people_by_official_rank(groups.get("local", []))
            groups["invited"] = self._sort_people_by_official_rank(groups.get("invited", []))

        return result

    def _get_all_selected_people_for_agenda(self):
        """
        Берет выбранных людей из 'Органы МСУ' и 'Приглашенные',
        подтягивает должности и готовит их для блока
        'Выбранные участники совещания'.
        """
        local_people = self._get_participants_with_positions_from_list(
            self.temp_selected_participants
        )

        invited_people = self._get_invited_with_positions_from_list(
            self.temp_invited_participants
        )

        local_people = self._sort_local_participants(local_people)
        invited_people = self._sort_invited_participants(invited_people)

        return local_people + invited_people

    def refresh_agenda_speaker_rows(self):
        """
        Обновляет блок 'Выбранные участники совещания'.
        Если ранее у человека уже стояла/снята галочка,
        состояние галочки сохраняется.
        """
        old_checked = {}
        old_question_numbers = {}

        for row in getattr(self, "agenda_speaker_rows", []):
            key = (
                row["fio"].get().strip(),
                row["position"].get().strip()
            )
            old_checked[key] = row["checked"].get()
            old_question_numbers[key] = row.get(
                "question_number",
                tk.StringVar(value="1")
            ).get().strip()

        self.agenda_speaker_rows = []

        people = self._get_all_selected_people_for_agenda()

        for person in people:
            fio = person.get("fio", "").strip()
            position = person.get("dolzhnost", "").strip()
            source = person.get("source", "local")

            if not fio:
                continue

            key = (fio, position)

            self.agenda_speaker_rows.append({
                "fio": tk.StringVar(value=fio),
                "position": tk.StringVar(value=position),
                "checked": tk.BooleanVar(value=old_checked.get(key, True)),
                "question_number": tk.StringVar(value=old_question_numbers.get(key, "1")),
                "source": source
            })

        if getattr(self, "tasks_mode", "") == "agenda" and hasattr(self, "render_agenda_speaker_rows"):
            self.render_agenda_speaker_rows()

    def _get_checked_agenda_speakers_for_doc(self):
        """
        Возвращает только тех участников, у кого стоит галочка
        в блоке 'Выбранные участники совещания'.
        Именно они попадут после 'Докладывают:' в документ 'Повестка'.
        """
        result = []

        for row in getattr(self, "agenda_speaker_rows", []):
            if not row["checked"].get():
                continue

            fio = row["fio"].get().strip()
            position = row["position"].get().strip()

            if not fio:
                continue

            result.append({
                "fio": fio,
                "dolzhnost": position or "Не указана",
                "question_number": self._normalize_agenda_question_number(
                    row.get("question_number", tk.StringVar(value="1")).get()
                ),
                "source": row.get("source", "local")
            })

        return result

    def generate_agenda_doc(self):
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from tkinter import filedialog, messagebox

        # В повестку попадают НЕ все выбранные в "Органы МСУ" / "Приглашенные",
        # а только те, у кого стоит галочка в блоке
        # "Выбранные участники совещания".
        # В 4-м поле указывается номер вопроса повестки.
        self.refresh_agenda_speaker_rows()
        agenda_people_by_question = self._get_checked_agenda_people_by_question_for_doc()
        agenda_questions = self.get_agenda_questions_for_doc()

        out_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"Повестка_{self.ent_date.get().replace('.', '_')}.docx"
        )

        if not out_path:
            return

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "PT Astra Serif"
        style.font.size = Pt(13)

        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)

        def set_font(paragraph, size=13, bold=False):
            for run in paragraph.runs:
                run.font.name = "PT Astra Serif"
                run.font.size = Pt(size)
                run.bold = bold

        def set_cell_width(cell, width_cm):
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()

            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)

            tc_w.set(qn("w:w"), str(int(width_cm * 567)))
            tc_w.set(qn("w:type"), "dxa")

        def set_table_width(table, width_cm):
            tbl = table._tbl
            tbl_pr = tbl.tblPr

            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)

            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(int(width_cm * 567)))

        def set_table_fixed_layout(table):
            tbl_pr = table._tbl.tblPr
            tbl_layout = tbl_pr.find(qn("w:tblLayout"))

            if tbl_layout is None:
                tbl_layout = OxmlElement("w:tblLayout")
                tbl_pr.append(tbl_layout)

            tbl_layout.set(qn("w:type"), "fixed")

        def remove_table_borders(table):
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._tc.get_or_add_tcPr()

                    old_borders = tc_pr.find(qn("w:tcBorders"))
                    if old_borders is not None:
                        tc_pr.remove(old_borders)

                    borders = OxmlElement("w:tcBorders")

                    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "nil")
                        borders.append(border)

                    tc_pr.append(borders)

        def justify_cell_text(cell):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                set_font(paragraph, 13)

        def add_people_block(label_text, people):
            if not people:
                return

            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.allow_autofit = False
            set_table_fixed_layout(table)
            set_table_width(table, 17.3)
            remove_table_borders(table)

            cell1 = table.rows[0].cells[0]
            cell2 = table.rows[0].cells[1]

            set_cell_width(cell1, 4.4)
            set_cell_width(cell2, 12.9)

            cell1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            cell1.text = ""
            p_cell1 = cell1.paragraphs[0]
            p_cell1.paragraph_format.line_spacing = 1
            p_cell1.paragraph_format.space_before = Pt(0)
            p_cell1.paragraph_format.space_after = Pt(0)
            p_cell1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_cell1.add_run(label_text)
            set_font(p_cell1, 13)

            cell2.text = ""

            for i, participant in enumerate(people):
                para = cell2.paragraphs[0] if i == 0 else cell2.add_paragraph()
                para.paragraph_format.line_spacing = 1
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.add_run(f"{participant['fio']}     -     {participant['dolzhnost']}")
                set_font(para, 13)

                if i != len(people) - 1:
                    empty_p = cell2.add_paragraph()
                    empty_p.paragraph_format.line_spacing = 1
                    empty_p.paragraph_format.space_before = Pt(0)
                    empty_p.paragraph_format.space_after = Pt(0)

            justify_cell_text(cell2)

        # ===== УТВЕРЖДАЮ =====
        approve = doc.add_paragraph()
        approve.paragraph_format.left_indent = Cm(10)
        approve.paragraph_format.line_spacing = 1
        approve.add_run("УТВЕРЖДАЮ\n")
        approve.add_run("Заместитель главы администрации\n")
        approve.add_run("города, руководитель аппарата\n\n")
        approve.add_run(f"«___» ______________ {datetime.now().year} г.")
        set_font(approve, 13)

        doc.add_paragraph()
        doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing = 1
        title.paragraph_format.space_after = Pt(0)
        title.paragraph_format.space_before = Pt(0)
        title.add_run("ПОВЕСТКА")
        set_font(title, 13)

        topic = doc.add_paragraph()
        topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        topic.paragraph_format.line_spacing = 1
        topic.paragraph_format.space_after = Pt(0)
        topic.paragraph_format.space_before = Pt(0)
        topic.add_run(f"совещание по вопросу {self.ent_title.get().strip()}")
        set_font(topic, 13)

        doc.add_paragraph()
        doc.add_paragraph()

        info = doc.add_paragraph()
        info.paragraph_format.left_indent = Cm(10.7)
        info.paragraph_format.line_spacing = 1
        info.add_run(f"{self.ent_date.get()}\n")
        info.add_run(f"{self.ent_time.get()} час.\n")
        info.add_run(f"ул. Гоголя, 48, каб. {self.cabinet_number.get().strip() or '213'}")
        set_font(info, 13)

        doc.add_paragraph()
        doc.add_paragraph()

        # ===== ВОПРОСЫ ПОВЕСТКИ + ДОКЛАДЫВАЮТ / ИНФОРМИРУЮТ =====
        for index, question_text in enumerate(agenda_questions, start=1):
            p_question = doc.add_paragraph()
            p_question.paragraph_format.first_line_indent = Cm(1)
            p_question.paragraph_format.line_spacing = 1
            p_question.paragraph_format.space_before = Pt(0)
            p_question.paragraph_format.space_after = Pt(0)
            p_question.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_question.add_run(f"{index}.  ")
            p_question.add_run(question_text)
            set_font(p_question, 13)

            people_for_question = agenda_people_by_question.get(index, {})
            local_speakers = people_for_question.get("local", [])
            invited_speakers = people_for_question.get("invited", [])

            if local_speakers:
                add_people_block(
                    "Докладывает:" if len(local_speakers) == 1 else "Докладывают:",
                    local_speakers
                )

            if invited_speakers:
                add_people_block(
                    "Информирует:" if len(invited_speakers) == 1 else "Информируют:",
                    invited_speakers
                )

            doc.add_paragraph()

        doc.add_paragraph()

        # ===== ПОДПИСЬ =====
        sign_table = doc.add_table(rows=1, cols=2)
        sign_table.autofit = False
        sign_table.allow_autofit = False
        sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(sign_table)
        set_table_width(sign_table, 17.3)
        remove_table_borders(sign_table)

        left = sign_table.rows[0].cells[0]
        right = sign_table.rows[0].cells[1]

        set_cell_width(left, 12.8)
        set_cell_width(right, 4.5)

        left.text = ""
        right.text = ""

        sign_position = self.agenda_sign_position.get().strip() or "Председатель правового комитета"
        sign_fio = self.agenda_sign_fio.get().strip() or "О.И. Насыров"

        position_words = sign_position.split()

        if len(position_words) > 3:
            first_line = " ".join(position_words[:2])
            second_line = " ".join(position_words[2:])

            p_left_1 = left.paragraphs[0]
            p_left_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left_1.paragraph_format.line_spacing = 1
            p_left_1.add_run(first_line)
            set_font(p_left_1, 13)

            p_left_2 = left.add_paragraph()
            p_left_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left_2.paragraph_format.line_spacing = 1
            p_left_2.add_run(second_line)
            set_font(p_left_2, 13)

            p_right_1 = right.paragraphs[0]
            p_right_1.paragraph_format.line_spacing = 1
            p_right_1.add_run("")
            set_font(p_right_1, 13)

            p_right_2 = right.add_paragraph()
            p_right_2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right_2.paragraph_format.line_spacing = 1
            p_right_2.add_run(sign_fio)
            set_font(p_right_2, 13)

        else:
            p_left = left.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left.paragraph_format.line_spacing = 1
            p_left.add_run(sign_position)
            set_font(p_left, 13)

            p_right = right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.line_spacing = 1
            p_right.add_run(sign_fio)
            set_font(p_right, 13)

        saved_path = self.safe_save_docx(doc, out_path)

        if not saved_path:
            return

        messagebox.showinfo("Успех", f"✅ Повестка создана:\n{saved_path}")

    def get_previous_workday(self, date_text):
        """
        Возвращает рабочий день перед датой совещания.
        Если совещание в понедельник, вернет пятницу.
        """
        from datetime import datetime, timedelta

        try:
            date_obj = datetime.strptime(date_text.strip(), "%d.%m.%Y")
        except Exception:
            return date_text

        date_obj -= timedelta(days=1)

        while date_obj.weekday() >= 5:  # 5 — суббота, 6 — воскресенье
            date_obj -= timedelta(days=1)

        return date_obj.strftime("%d.%m.%Y")

    def generate_telegram_doc(self):
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from tkinter import filedialog, messagebox

        selected = self.temp_selected_participants
        invited_selected = self.temp_invited_participants

        if not selected and not invited_selected:
            messagebox.showwarning("Внимание", "Выберите хотя бы одного сотрудника или приглашенного.")
            return

        all_emps = self._load_local_people_records()

        if not all_emps:
            all_emps = get_all_employees()

        try:
            all_invited_people = self._load_invited_people_records()

            if not all_invited_people:
                all_invited_people = invited_get_all()
        except Exception:
            all_invited_people = []

        title_text = self.ent_title.get().strip()
        telegram_number = self.telegram_number.get().strip()

        if not title_text:
            messagebox.showwarning("Внимание", "Заполните поле 'По вопросу:'.")
            return

        if not telegram_number:
            messagebox.showwarning("Внимание", "Введите номер телефонограммы.")
            return

        deadline_date = self.get_previous_workday(self.ent_date.get())

        def short_fio(full_name):
            parts = full_name.strip().split()
            if len(parts) >= 3:
                return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
            return full_name

        def first_patronymic(full_name):
            parts = full_name.strip().split()
            if len(parts) >= 3:
                return f"{parts[1]} {parts[2]}"
            return full_name

        def set_font(paragraph, size=14, bold=False):
            for run in paragraph.runs:
                run.font.name = "PT Astra Serif"
                run.font.size = Pt(size)
                run.bold = bold

        def set_cell_width(cell, width_cm):
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()

            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)

            tc_w.set(qn("w:w"), str(int(width_cm * 567)))
            tc_w.set(qn("w:type"), "dxa")

        def set_table_width(table, width_cm):
            tbl = table._tbl
            tbl_pr = tbl.tblPr

            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)

            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(int(width_cm * 567)))

        def set_table_fixed_layout(table):
            tbl_pr = table._tbl.tblPr
            tbl_layout = tbl_pr.find(qn("w:tblLayout"))

            if tbl_layout is None:
                tbl_layout = OxmlElement("w:tblLayout")
                tbl_pr.append(tbl_layout)

            tbl_layout.set(qn("w:type"), "fixed")

        def set_row_height(row, height_twips):
            tr_pr = row._tr.get_or_add_trPr()

            tr_height = tr_pr.find(qn("w:trHeight"))
            if tr_height is None:
                tr_height = OxmlElement("w:trHeight")
                tr_pr.append(tr_height)

            tr_height.set(qn("w:val"), str(height_twips))
            tr_height.set(qn("w:hRule"), "exact")

        def remove_table_borders(table):
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._tc.get_or_add_tcPr()

                    old_borders = tc_pr.find(qn("w:tcBorders"))
                    if old_borders is not None:
                        tc_pr.remove(old_borders)

                    borders = OxmlElement("w:tcBorders")

                    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "nil")
                        borders.append(border)

                    tc_pr.append(borders)

        def add_double_line(doc):
            # Жирная линия
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(0)
            p1.paragraph_format.line_spacing = 0.1

            p_pr = p1._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "18")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")
            borders.append(bottom)
            p_pr.append(borders)

            # Тонкая линия максимально близко
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.line_spacing = 0.1

            p_pr = p2._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "000000")
            borders.append(bottom)
            p_pr.append(borders)

        def clear_paragraph_spacing(paragraph):
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

        def add_empty_paragraph(doc):
            p = doc.add_paragraph()
            clear_paragraph_spacing(p)
            return p

        def decline_fio_genitive(full_name):
            """ФИО в родительном падеже: Курышин Андрей Александрович -> Курышина А.А."""
            parts = str(full_name or "").strip().split()

            if len(parts) < 3:
                return full_name

            surname, name, patronymic = parts[0], parts[1], parts[2]

            def surname_rp(s):
                if s.endswith(("ов", "ев", "ёв", "ин", "ын")):
                    return s + "а"
                if s.endswith(("ский", "цкий")):
                    return s[:-2] + "ого"
                if s.endswith(("ова", "ева", "ёва", "ина", "ына")):
                    return s[:-1] + "ой"
                if s.endswith("ая"):
                    return s[:-2] + "ой"
                if s.endswith("яя"):
                    return s[:-2] + "ей"
                return s

            def name_rp(s):
                irregular = {
                    "Павел": "Павла",
                    "Лев": "Льва",
                    "Пётр": "Петра",
                    "Петр": "Петра",
                    "Илья": "Ильи",
                    "Никита": "Никиты",
                    "Фома": "Фомы",
                    "Кузьма": "Кузьмы",
                }

                if s in irregular:
                    return irregular[s]

                if s.endswith("й"):
                    return s[:-1] + "я"
                if s.endswith("ь"):
                    return s[:-1] + "я"
                if s.endswith("а"):
                    return s[:-1] + "ы"
                if s.endswith("я"):
                    return s[:-1] + "и"
                return s + "а"

            def patronymic_rp(s):
                if s.endswith("ич"):
                    return s + "а"
                if s.endswith("на"):
                    return s[:-1] + "ы"
                return s

            return f"{surname_rp(surname)} {name_rp(name)} {patronymic_rp(patronymic)}"

        def decline_fio_dative(full_name):
            """ФИО в дательном падеже: Артемов Александр Владимирович -> Артемову А.В."""
            parts = str(full_name or "").strip().split()

            if len(parts) < 3:
                return full_name

            surname, name, patronymic = parts[0], parts[1], parts[2]

            def surname_dp(s):
                if s.endswith(("ов", "ев", "ёв", "ин", "ын")):
                    return s + "у"
                if s.endswith(("ский", "цкий")):
                    return s[:-2] + "ому"
                if s.endswith(("ова", "ева", "ёва", "ина", "ына")):
                    return s[:-1] + "ой"
                if s.endswith("ая"):
                    return s[:-2] + "ой"
                if s.endswith("яя"):
                    return s[:-2] + "ей"
                return s

            def name_dp(s):
                if s.endswith("й"):
                    return s[:-1] + "ю"
                if s.endswith("ь"):
                    return s[:-1] + "ю"
                if s.endswith("а"):
                    return s[:-1] + "е"
                if s.endswith("я"):
                    return s[:-1] + "е"
                return s + "у"

            def patronymic_dp(s):
                if s.endswith("ич"):
                    return s + "у"
                if s.endswith("на"):
                    return s[:-1] + "е"
                return s

            return f"{surname_dp(surname)} {name_dp(name)} {patronymic_dp(patronymic)}"

        def initials_surname_dative(full_name):
            """
            Артемов Александр Владимирович -> А.В. Артемову.
            Должность и ФИО адресата в телефонограмме должны быть в дательном падеже,
            при этом сначала идут инициалы, затем фамилия.
            """
            parts = str(full_name or "").strip().split()

            if len(parts) < 3:
                return full_name

            surname, name, patronymic = parts[0], parts[1], parts[2]
            declined_parts = decline_fio_dative(full_name).split()
            declined_surname = declined_parts[0] if declined_parts else surname

            return f"{name[0]}.{patronymic[0]}. {declined_surname}"

        def make_email_wrappable(email):
            """
            Возвращает электронную почту как цельный фрагмент.
            Если она не помещается в конце строки, Word переносит ее целиком
            на следующую строку, а не разбивает по точкам или дефисам.
            """
            email = str(email or "").strip()
            return email.replace("-", "\u2011")

        def normalize_for_match(value):
            value = str(value or "").strip().lower().replace("ё", "е")
            return " ".join(value.split())

        def invited_person_name(person):
            return (
                person.get("full_name")
                or person.get("fio")
                or person.get("name")
                or ""
            ).strip()

        def invited_person_position(person):
            return (
                person.get("position")
                or person.get("dolzhnost")
                or person.get("job_title")
                or ""
            ).strip()

        def invited_person_leadership(person):
            return (
                person.get("leadership")
                or person.get("rukovodstvo")
                or person.get("supervisor")
                or person.get("boss")
                or ""
            ).strip()

        def get_selected_invited_records():
            records = []

            for selected_name in invited_selected:
                person = next(
                    (
                        p for p in all_invited_people
                        if invited_person_name(p) == str(selected_name).strip()
                    ),
                    None
                )

                if not person:
                    continue

                records.append({
                    "full_name": invited_person_name(person),
                    "position": invited_person_position(person),
                    "leadership": invited_person_leadership(person),
                })

            return records

        selected_invited_records = get_selected_invited_records()

        def is_invited_deputy_record(person):
            """Приглашенный считается депутатом, если в должности есть слово 'депутат'."""
            return "депутат" in normalize_for_match(person.get("position", ""))

        def find_employee_by_position(position):
            position_key = normalize_for_match(position)

            if not position_key:
                return None

            return next(
                (
                    emp for emp in all_emps
                    if normalize_for_match(emp.get("position", "")) == position_key
                ),
                None
            )

        def find_invited_by_position(position):
            position_key = normalize_for_match(position)

            if not position_key:
                return None

            return next(
                (
                    person for person in all_invited_people
                    if normalize_for_match(invited_person_position(person)) == position_key
                ),
                None
            )

        def build_invited_deputy_groups():
            """
            Формирует отдельные страницы телефонограммы для выбранных приглашенных,
            у которых в должности есть слово 'депутат' и заполнено руководство.

            Адресат страницы — руководитель депутата из того же раздела
            'Список приглашенных': должность берется из leadership, ФИО — из записи,
            где position совпадает с leadership.
            """
            deputy_records = [
                person for person in selected_invited_records
                if is_invited_deputy_record(person) and person.get("leadership", "").strip()
            ]

            groups_by_leadership = {}

            for person in deputy_records:
                leadership = person.get("leadership", "").strip()
                key = normalize_for_match(leadership)

                groups_by_leadership.setdefault(key, {
                    "leadership": leadership,
                    "people": []
                })

                groups_by_leadership[key]["people"].append(person)

            result = []

            for group_data in groups_by_leadership.values():
                leadership = group_data.get("leadership", "").strip()
                leader_person = find_invited_by_position(leadership)

                if leader_person:
                    address_emp = {
                        "position": invited_person_position(leader_person) or leadership,
                        "full_name": invited_person_name(leader_person)
                    }
                else:
                    # Если руководитель указан должностью, но его запись с ФИО еще не заведена,
                    # страницу все равно формируем: должность будет правильной, ФИО останется пустым.
                    address_emp = {
                        "position": leadership,
                        "full_name": ""
                    }

                result.append({
                    "deputy": address_emp,
                    "employees": [],
                    "invited_override": group_data["people"],
                    "omit_report_paragraphs": True
                })

            return result

        def build_invited_self_groups():
            """
            Формирует отдельные страницы для выбранных приглашенных без руководства.
            Такие телефонограммы адресуются самому приглашенному: его должность и ФИО
            ставятся в правый адресный блок в дательном падеже.
            """
            result = []

            for person in selected_invited_records:
                if person.get("leadership", "").strip():
                    continue

                result.append({
                    "deputy": {
                        "position": person.get("position", ""),
                        "full_name": person.get("full_name", "")
                    },
                    "employees": [],
                    "invited_override": None,
                    "invited_self_page": True
                })

            return result

        def decline_position_genitive(position):
            """Должность в родительном падеже"""
            text = str(position or "").strip()

            replacements = {
                                "Глава администрации района":
                    "Главы администрации района",

                "Глава администрации Индустриального района":
                    "Главы администрации Индустриального района",

                "Глава администрации Октябрьского района":
                    "Главы администрации Октябрьского района",

                "Глава администрации Железнодорожного района":
                    "Главы администрации Железнодорожного района",

                "Глава администрации Ленинского района":
                    "Главы администрации Ленинского района",

                "Глава администрации Центрального района":
                    "Главы администрации Центрального района",

                "Заместитель главы администрации района":
                    "Заместителя главы администрации района",

                "Заместитель главы администрации Индустриального района":
                    "Заместителя главы администрации Индустриального района",

                "Заместитель главы администрации Октябрьского района":
                    "Заместителя главы администрации Октябрьского района",

                "Заместитель главы администрации Железнодорожного района":
                    "Заместителя главы администрации Железнодорожного района",

                "Заместитель главы администрации Ленинского района":
                    "Заместителя главы администрации Ленинского района",

                "Заместитель главы администрации Центрального района":
                    "Заместителя главы администрации Центрального района",

                "Заместитель председателя комитета":
                    "Заместителя председателя комитета",

                "Заведующий отделом":
                    "Заведующего отделом",
            }

            return replacements.get(text, text)

        def decline_position_dative(position):
            """Должность в дательном падеже"""
            text = str(position or "").strip()

            replacements = {
                "Первый заместитель главы администрации города":
                    "Первому заместителю главы администрации города",

                "Заместитель главы администрации города":
                    "Заместителю главы администрации города",

                "Заместитель главы администрации города, руководитель аппарата":
                    "Заместителю главы администрации города, руководителю аппарата",

                "Заместитель главы администрации города по экономической политике":
                    "Заместителю главы администрации города по экономической политике",

                "Заместитель главы администрации города по защите населения и информации":
                    "Заместителю главы администрации города по защите населения и информации",

                "Заместитель главы администрации города по городскому хозяйству":
                    "Заместителю главы администрации города по городскому хозяйству",

                "Заместитель главы администрации города по социальной политике":
                    "Заместителю главы администрации города по социальной политике",

                "Заместитель главы администрации города по дорожному хозяйству и транспорту":
                    "Заместителю главы администрации города по дорожному хозяйству и транспорту",

                "Глава администрации района":
                    "Главе администрации района",

                "Глава администрации Индустриального района":
                    "Главе администрации Индустриального района",

                "Глава администрации Октябрьского района":
                    "Главе администрации Октябрьского района",

                "Глава администрации Железнодорожного района":
                    "Главе администрации Железнодорожного района",

                "Глава администрации Ленинского района":
                    "Главе администрации Ленинского района",

                "Глава администрации Центрального района":
                    "Главе администрации Центрального района",

                "Заместитель главы администрации района":
                    "Заместителю главы администрации района",

                "Заместитель главы администрации Индустриального района":
                    "Заместителю главы администрации Индустриального района",

                "Заместитель главы администрации Октябрьского района":
                    "Заместителю главы администрации Октябрьского района",

                "Заместитель главы администрации Железнодорожного района":
                    "Заместителю главы администрации Железнодорожного района",

                "Заместитель главы администрации Ленинского района":
                    "Заместителю главы администрации Ленинского района",

                "Заместитель главы администрации Центрального района":
                    "Заместителю главы администрации Центрального района",

                "Заместитель председателя комитета":
                    "Заместителю председателя комитета",

                "Заведующий отделом":
                    "Заведующему отделом",
            }

            if text in replacements:
                return replacements[text]

            lowered = text.lower().replace("ё", "е")

            if lowered.startswith("председатель "):
                return "Председателю " + text[len("Председатель "):]

            if lowered.startswith("депутат "):
                return "Депутату " + text[len("Депутат "):]

            if lowered.startswith("директор "):
                return "Директору " + text[len("Директор "):]

            return text

        def is_female_full_name(full_name):
            parts = str(full_name or "").strip().split()

            # Женщину определяем по отчеству или фамилии
            if len(parts) >= 3:
                patronymic = parts[2]
                if patronymic.endswith(("на", "зы")):
                    return True

            if len(parts) >= 1:
                surname = parts[0]
                if surname.endswith(("ова", "ева", "ёва", "ина", "ына", "ая", "яя")):
                    return True

            return False

        def decline_surname_for_invite(full_name):
            """
            Фамилия для фразы 'Приглашаем Вас, а также ...'

            Мужчина:
            Курышин Андрей Александрович -> Курышина А.А.

            Женщина:
            Королева Татьяна Николаевна -> Королеву Т.Н.
            """
            parts = str(full_name or "").strip().split()

            if len(parts) < 3:
                return full_name

            surname, name, patronymic = parts[0], parts[1], parts[2]
            female = is_female_full_name(full_name)

            if female:
                if surname.endswith(("ова", "ева", "ёва", "ина", "ына")):
                    declined_surname = surname[:-1] + "у"
                elif surname.endswith("ая"):
                    declined_surname = surname[:-2] + "ую"
                elif surname.endswith("яя"):
                    declined_surname = surname[:-2] + "юю"
                else:
                    declined_surname = surname
            else:
                if surname.endswith(("ов", "ев", "ёв", "ин", "ын")):
                    declined_surname = surname + "а"
                elif surname.endswith(("ский", "цкий")):
                    declined_surname = surname[:-2] + "ого"
                else:
                    declined_surname = surname

            return f"{declined_surname} {name[0]}.{patronymic[0]}."

        def decline_position_genitive(position):
            """Должность в родительном падеже"""
            text = str(position or "").strip()

            replacements = {
                "Первый заместитель главы администрации города":
                    "Первого заместителя главы администрации города",

                "Заместитель главы администрации города, руководитель аппарата":
                    "Заместителя главы администрации города, руководителя аппарата",

                "Заместитель главы администрации города по экономической политике":
                    "Заместителя главы администрации города по экономической политике",

                "Заместитель главы администрации города по защите населения и информации":
                    "Заместителя главы администрации города по защите населения и информации",

                "Заместитель главы администрации города по городскому хозяйству":
                    "Заместителя главы администрации города по городскому хозяйству",

                "Заместитель главы администрации города по социальной политике":
                    "Заместителя главы администрации города по социальной политике",

                "Заместитель главы администрации города по дорожному хозяйству и транспорту":
                    "Заместителя главы администрации города по дорожному хозяйству и транспорту",

                "Председатель комитета":
                    "Председателя комитета",

                "Председатель правового комитета":
                    "Председателя правового комитета",

                "Начальник управления":
                    "Начальника управления",

                "Начальник отдела":
                    "Начальника отдела",

                "Управляющий делами администрации, председатель комитета":
                    "Управляющего делами администрации, председателя комитета",
            }

            if text.startswith("Депутат "):
                return "Депутата " + text[len("Депутат "):]

            return replacements.get(text, text)

        def format_employee_for_invite(emp):
            """
            Формирует строку для текста телефонограммы:

            Было:
            Председатель комитета Курышин А.А.

            Стало:
            Председателя комитета Курышина А.А.

            Для женщины:
            Председателя комитета Королеву Т.Н.
            """
            position = emp.get("position", "").strip()
            full_name = emp.get("full_name", "").strip()

            position_rp = decline_position_genitive(position)
            fio_needed_case = decline_surname_for_invite(full_name)

            if position_rp:
                return f"{position_rp} {fio_needed_case}"

            return fio_needed_case

        def build_grouped_by_deputy():
            """
            Формирует группы телефонограмм для выбранных из "Органы МСУ".

            Новый источник — таблица meeting_participants/category='msu_ip',
            которая заполняется в настройках "Участники совещаний".
            Если у выбранного участника указано поле leadership, его телефонограмма
            адресуется записи, у которой position совпадает с этим leadership.
            Если leadership пустое, участник сам становится адресатом страницы.
            """
            groups = {}

            def person_name(person):
                return (
                    person.get("full_name")
                    or person.get("fio")
                    or person.get("name")
                    or ""
                ).strip()

            def person_position(person):
                return (
                    person.get("position")
                    or person.get("dolzhnost")
                    or person.get("job_title")
                    or ""
                ).strip()

            def person_leadership(person):
                return (
                    person.get("leadership")
                    or person.get("rukovodstvo")
                    or person.get("supervisor")
                    or person.get("boss")
                    or ""
                ).strip()

            def normalize(value):
                return normalize_for_match(value)

            def find_by_name(name):
                return next(
                    (
                        e for e in all_emps
                        if person_name(e) == str(name).strip()
                    ),
                    None
                )

            def find_by_position(position):
                position_key = normalize(position)

                if not position_key:
                    return None

                return next(
                    (
                        e for e in all_emps
                        if normalize(person_position(e)) == position_key
                    ),
                    None
                )

            for selected_name in selected:
                emp = find_by_name(selected_name)

                if not emp:
                    messagebox.showerror("Ошибка", f"Сотрудник не найден:\n{selected_name}")
                    continue

                emp_leadership = person_leadership(emp)

                # Новая схема из настроек meeting_participants.
                if emp_leadership or emp.get("source_table") == "meeting_participants":
                    if emp_leadership:
                        deputy = find_by_position(emp_leadership)

                        if not deputy:
                            messagebox.showerror(
                                "Ошибка",
                                f"Не найден руководитель для сотрудника:\n{person_name(emp)}\n\n"
                                f"В поле 'Руководство' указано:\n{emp_leadership}\n\n"
                                "Проверьте, что в настройках 'Список органов МСУ' есть запись "
                                "с такой должностью."
                            )
                            continue

                        deputy_key = f"mp_{deputy.get('id') or person_position(deputy)}"

                        groups.setdefault(deputy_key, {
                            "deputy": {
                                "id": deputy.get("id"),
                                "full_name": person_name(deputy),
                                "position": person_position(deputy),
                            },
                            "employees": []
                        })

                        groups[deputy_key]["employees"].append({
                            "id": emp.get("id"),
                            "full_name": person_name(emp),
                            "position": person_position(emp),
                        })
                    else:
                        deputy_key = f"mp_{emp.get('id') or person_position(emp)}"

                        groups.setdefault(deputy_key, {
                            "deputy": {
                                "id": emp.get("id"),
                                "full_name": person_name(emp),
                                "position": person_position(emp),
                            },
                            "employees": []
                        })

                    continue

                # Старый формат get_all_employees().
                if emp.get("level") == "deputy":
                    deputy_id = emp.get("id")
                    groups.setdefault(deputy_id, {
                        "deputy":emp,
                        "employees":[]
                    })
                    continue

                supervisor_id = emp.get("supervisor_id")
                deputy = next(
                    (e for e in all_emps if e.get("id") == supervisor_id),
                    None
                )

                if not deputy:
                    messagebox.showerror(
                        "Ошибка",
                        f"Не найден руководитель для сотрудника:\n{emp.get('full_name', '')}"
                    )
                    continue

                deputy_id = deputy.get("id")

                groups.setdefault(deputy_id, {
                    "deputy":deputy,
                    "employees":[]
                })

                groups[deputy_id]["employees"].append(emp)

            result_groups = list(groups.values())

            for group in result_groups:
                group["employees"] = self._sort_people_by_official_rank([
                    {
                        "fio": person_name(emp),
                        "dolzhnost": person_position(emp),
                        "employee": emp,
                        "source":"local"
                    }
                    for emp in group.get("employees", [])
                ])

                group["employees"] = [
                    person.get("employee")
                    for person in group["employees"]
                    if person.get("employee")
                ]

            result_groups = sorted(
                result_groups,
                key=lambda group:(
                    self._official_position_rank(group["deputy"].get("position", "")),
                    self._normalize_text_for_sort(group["deputy"].get("full_name", ""))
                )
            )

            return result_groups

        out_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"Телефонограммы_{self.ent_date.get().replace('.', '_')}.docx"
        )

        if not out_path:
            return

        doc = Document()

        def force_telegram_page_setup(section):
            """
            Жестко задает параметры листа телефонограммы через свойства
            python-docx и напрямую через XML.

            Это нужно, чтобы Word не оставлял старые поля 1,5 см слева
            и 0,5 см справа: в XML поля страницы хранятся в twips.
            2 см = 1134 twips, 1 см = 567 twips.
            """
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(1)
            section.gutter = Cm(0)

            sect_pr = section._sectPr
            pg_mar = sect_pr.find(qn("w:pgMar"))

            if pg_mar is None:
                pg_mar = OxmlElement("w:pgMar")
                sect_pr.append(pg_mar)

            pg_mar.set(qn("w:top"), "1134")
            pg_mar.set(qn("w:bottom"), "1134")
            pg_mar.set(qn("w:left"), "1134")
            pg_mar.set(qn("w:right"), "567")
            pg_mar.set(qn("w:gutter"), "0")

            pg_sz = sect_pr.find(qn("w:pgSz"))

            if pg_sz is None:
                pg_sz = OxmlElement("w:pgSz")
                sect_pr.append(pg_sz)

            # A4: 21 x 29.7 см = 11906 x 16838 twips.
            pg_sz.set(qn("w:w"), "11906")
            pg_sz.set(qn("w:h"), "16838")

        section = doc.sections[0]
        force_telegram_page_setup(section)

        style = doc.styles["Normal"]
        style.font.name = "PT Astra Serif"
        style.font.size = Pt(14)

        def add_transfer_text_at_page_bottom():
            """
            Добавляет блок "Передала" как основной текст документа, а не как
            нижний колонтитул и не как таблицу.

            Текст закрепляется у нижней границы рабочей области листа. Последняя
            строка с телефоном находится у нижнего поля без дополнительных
            пустых абзацев после нее.
            """
            transfer_fio = self.transfer_fio.get().strip()
            transfer_phone = self.transfer_phone.get().strip()

            lines = ["Передала:"]
            if transfer_fio:
                lines.append(transfer_fio)
            if transfer_phone:
                lines.append(transfer_phone)

            p = doc.add_paragraph()
            clear_paragraph_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run("\n".join(lines))
            set_font(p, size=10)

            # Размещаем именно абзац основного текста в фиксированной области
            # у нижнего поля страницы. Это не footer и не таблица.
            p_pr = p._p.get_or_add_pPr()
            frame_pr = OxmlElement("w:framePr")
            frame_pr.set(qn("w:w"), "9000")
            frame_pr.set(qn("w:h"), "900")
            frame_pr.set(qn("w:x"), "1134")
            # A4: 29.7 см. Нижнее поле 2 см. Высота блока около 0.9 см.
            # Поэтому верх рамки ставим примерно на 26.8 см от верха листа,
            # чтобы последняя строка телефона была у нижнего поля без пустых абзацев.
            frame_pr.set(qn("w:y"), "14350")
            frame_pr.set(qn("w:hAnchor"), "page")
            frame_pr.set(qn("w:vAnchor"), "page")
            frame_pr.set(qn("w:wrap"), "none")
            frame_pr.set(qn("w:hRule"), "atLeast")
            p_pr.append(frame_pr)

        def add_telegram_page(group, is_first_page=True):
            if not is_first_page:
                doc.add_page_break()

            address_emp = group["deputy"]
            group_employees = group["employees"]

            address_position = address_emp.get("position", "").strip()
            address_fio = address_emp.get("full_name", "").strip()

            address_position_dative = decline_position_dative(address_position)
            address_fio_dative_short = initials_surname_dative(address_fio) if address_fio else ""
            omit_report_paragraphs = bool(group.get("omit_report_paragraphs"))

            invited_override = group.get("invited_override")

            if invited_override is not None:
                invited_for_this_leadership = invited_override
            else:
                invited_for_this_leadership = [
                    person for person in selected_invited_records
                    if person.get("leadership")
                    and not is_invited_deputy_record(person)
                    and normalize_for_match(person.get("leadership")) == normalize_for_match(address_position)
                ]

            if group_employees:
                employees_text = ", ".join(
                    format_employee_for_invite(emp)
                    for emp in group_employees
                )
                add_text = f", а также {employees_text},"
            else:
                add_text = ""

            # ===== ВЕРХНИЙ ЗАГОЛОВОК =====
            header_table = doc.add_table(rows=1, cols=1)
            header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            header_table.autofit = False
            header_table.allow_autofit = False
            set_table_fixed_layout(header_table)
            set_table_width(header_table, 16.8)
            remove_table_borders(header_table)

            header_cell = header_table.rows[0].cells[0]
            set_cell_width(header_cell, 16.8)

            p = header_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_paragraph_spacing(p)
            p.add_run("Заместитель главы администрации города, руководитель\n")
            p.add_run("аппарата")
            set_font(p, size=14, bold=True)

            # Две линии под заголовком
            add_double_line(doc)

            # 1 абзац между двойной линией и надписью "ТЕЛЕФОНОГРАММА ..."
            add_empty_paragraph(doc)

            # ===== НОМЕР ТЕЛЕФОНОГРАММЫ =====
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_paragraph_spacing(p)

            number_text = telegram_number
            if not number_text.startswith("№"):
                number_text = f"№ {number_text}"

            p.add_run(f"ТЕЛЕФОНОГРАММА {number_text} от __________")
            set_font(p, size=14, bold=True)

            add_empty_paragraph(doc)

            # ===== АДРЕСАТ =====

            addr_table = doc.add_table(rows=2, cols=3)
            addr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            addr_table.autofit = False
            addr_table.allow_autofit = False
            set_table_fixed_layout(addr_table)
            set_table_width(addr_table, 17.2)

            col_widths = [10.5, 0.5, 6.2]

            for row in addr_table.rows:
                for idx, width in enumerate(col_widths):
                    set_cell_width(row.cells[idx], width)

            addr_table.rows[1].cells[0].merge(addr_table.rows[1].cells[1]).merge(addr_table.rows[1].cells[2])

            right_top = addr_table.rows[0].cells[2]
            appeal_cell = addr_table.rows[1].cells[0]

            right_top.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            appeal_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # ===== Адресат справа =====
            right_top.text = ""

            # 1 абзац перед строкой "Заместитель ..."
            p_empty_before = right_top.paragraphs[0]
            clear_paragraph_spacing(p_empty_before)

            # строка с должностью и ФИО
            p = right_top.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            clear_paragraph_spacing(p)
            p.add_run(f"{address_position_dative}\n")
            if address_fio_dative_short:
                p.add_run(address_fio_dative_short)
            set_font(p)

            # 2 абзаца после строки "Заместитель ..."
            p_empty_after_1 = right_top.add_paragraph()
            clear_paragraph_spacing(p_empty_after_1)

            p_empty_after_2 = right_top.add_paragraph()
            clear_paragraph_spacing(p_empty_after_2)

            # ===== Обращение по центру =====
            appeal_cell.text = ""

            p = appeal_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_paragraph_spacing(p)
            if address_fio:
                p.add_run(f"Уважаемый {first_patronymic(address_fio)}!")
            else:
                p.add_run("Уважаемый(ая)!")
            set_font(p)

            # 1 абзац после строки "Уважаемый ..."
            p_empty_after_appeal_1 = appeal_cell.add_paragraph()
            clear_paragraph_spacing(p_empty_after_appeal_1)

            remove_table_borders(addr_table)

            # ===== ОСНОВНОЙ ТЕКСТ =====
            if invited_for_this_leadership:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                clear_paragraph_spacing(p)
                p.add_run(
                    f"{self.ent_date.get()} в {self.ent_time.get()} час. "
                    f"по адресу: ул.Гоголя, 48, кабинет {self.cabinet_number.get().strip() or '213'} "
                    f"состоится совещание по вопросу {title_text}."
                )
                set_font(p)

                for invited_person in invited_for_this_leadership:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(1.25)
                    clear_paragraph_spacing(p)
                    invited_position_genitive = decline_position_genitive(
                        invited_person.get('position', '').strip()
                    )
                    invited_fio_genitive = decline_fio_genitive(
                        invited_person.get('full_name', '').strip()
                    )

                    p.add_run(
                        f"Просим направить для участия "
                        f"{invited_position_genitive} "
                        f"{invited_fio_genitive}."
                    )
                    set_font(p)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                clear_paragraph_spacing(p)
                p.add_run(
                    f"Приглашаем Вас{add_text} принять участие в совещании по вопросу "
                    f"{title_text}, которое состоится {self.ent_date.get()} в {self.ent_time.get()} час. "
                    f"по адресу: ул. Гоголя, 48, кабинет {self.cabinet_number.get().strip() or '213'}."
                )
                set_font(p)

            if not omit_report_paragraphs:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                clear_paragraph_spacing(p)
                p.add_run("В целях конструктивной работы просим быть готовыми к докладу согласно повестке.")
                set_font(p)

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                clear_paragraph_spacing(p)

                email = self.transfer_email.get().strip() or "fedorova-en@barnaul-adm.ru"
                p.add_run("Доклады просим направить на адрес электронной почты: ")
                p.add_run(make_email_wrappable(email))
                p.add_run(f" в срок до {deadline_date}.")
                set_font(p)

            # Между основным текстом и подписью О.А. Финк — ровно 2 абзаца.
            add_empty_paragraph(doc)
            add_empty_paragraph(doc)

            # ===== Подпись О.А. Финк =====
            sign_fink = doc.add_paragraph()
            sign_fink.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            clear_paragraph_spacing(sign_fink)
            sign_fink.add_run(self.telegram_sign_fio.get().strip() or "О.А. Финк")
            set_font(sign_fink)

            # Блок "Передала" должен быть основным текстом документа, а не колонтитулом.
            add_transfer_text_at_page_bottom()
            return True

        groups = build_grouped_by_deputy()
        invited_deputy_groups = build_invited_deputy_groups()
        invited_self_groups = build_invited_self_groups()

        if not groups and not invited_deputy_groups and not invited_self_groups:
            messagebox.showerror("Ошибка", "Не удалось сформировать телефонограммы: сотрудники или приглашенные не найдены.")
            return

        success_count = 0

        # Сначала формируем страницы по выбранным участникам из раздела "Органы МСУ".
        for group in groups:
            ok = add_telegram_page(group, is_first_page=(success_count == 0))

            if ok:
                success_count += 1

        # Затем формируем страницы для приглашенных без руководства.
        # Такие страницы адресуются самому приглашенному.
        for group in invited_self_groups:
            ok = add_telegram_page(group, is_first_page=(success_count == 0))

            if ok:
                success_count += 1

        # Затем в том же документе, на следующих листах, формируем страницы
        # для приглашенных депутатов с руководством.
        for group in invited_deputy_groups:
            ok = add_telegram_page(group, is_first_page=(success_count == 0))

            if ok:
                success_count += 1

        if success_count == 0:
            messagebox.showerror("Ошибка", "Не удалось сформировать ни одной телефонограммы.")
            return

        # На всякий случай применяем одинаковые поля ко всем секциям документа
        # перед сохранением. Поля задаются и через python-docx, и напрямую
        # через XML, чтобы Word точно показывал: слева 2 см, справа 1 см.
        for _section in doc.sections:
            force_telegram_page_setup(_section)

        saved_path = self.safe_save_docx(doc, out_path)

        if not saved_path:
            return

        messagebox.showinfo(
            "Успех",
            f"✅ Создан документ с телефонограммами: {success_count}\n{saved_path}"
        )

    def _short_fio(self, full_name):
        """
        Иванов Иван Иванович -> Иванов И.И.
        О.А. Финк остается как есть.
        """
        full_name = str(full_name or "").strip()
        parts = full_name.split()

        if len(parts) >= 3:
            return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."

        return full_name

    def _split_fio_for_protocol_table(self, full_name):
        parts = str(full_name or "").strip().split()

        if len(parts) >= 3:
            return parts[0], " ".join(parts[1:])

        if parts:
            return parts[0], ""

        return "", ""

    def _get_protocol_speakers_by_question(self):
        """
        Берет отмеченных участников из блока "Выбранные участники совещания"
        и группирует их по номеру вопроса повестки.
        Для протокола все они идут после "Докладывает/Докладывают"
        в именительном падеже, независимо от источника выбора.
        """
        result = {}

        for row in getattr(self, "agenda_speaker_rows", []):
            if not row["checked"].get():
                continue

            fio = row["fio"].get().strip()
            position = row["position"].get().strip()

            if not fio:
                continue

            question_number = self._normalize_agenda_question_number(
                row.get("question_number", tk.StringVar(value="1")).get()
            )

            result.setdefault(question_number, [])
            result[question_number].append({
                "fio": fio,
                "dolzhnost": position or "Не указана"
            })

        for question_number in list(result.keys()):
            result[question_number] = self._sort_people_by_official_rank(result[question_number])

        return result

    def generate_protocol_doc(self):
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from tkinter import filedialog, messagebox

        title_text = self.ent_title.get().strip()
        date_text = self.ent_date.get().strip()
        protocol_number = self.protocol_number.get().strip()
        chair_fio_raw = self.protocol_chair_fio.get().strip()
        chair_position = self.protocol_chair_position.get().strip()
        keeper_fio_raw = self.protocol_keeper.get().strip()

        if not title_text:
            messagebox.showwarning("Внимание", "Заполните поле 'По вопросу:'.")
            return

        if not date_text:
            messagebox.showwarning("Внимание", "Заполните поле 'Дата'.")
            return

        if not protocol_number:
            messagebox.showwarning("Внимание", "Заполните поле '№'.")
            return

        if not chair_fio_raw:
            messagebox.showwarning("Внимание", "Заполните поле 'Председательствующий'.")
            return

        if not chair_position:
            messagebox.showwarning("Внимание", "Заполните поле 'Должность'.")
            return

        if not keeper_fio_raw:
            messagebox.showwarning("Внимание", "Заполните поле 'Протокол вела'.")
            return

        def fio_surname_initials(full_name):
            """
            Финк Оксана Анатольевна -> Финк О.А.
            О.А. Финк -> Финк О.А.
            Финк О.А. -> Финк О.А.
            """
            full_name = str(full_name or "").strip()
            parts = full_name.split()

            if len(parts) >= 3:
                return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."

            if len(parts) == 2:
                first, second = parts[0], parts[1]

                if "." in first and "." not in second:
                    return f"{second} {first}"

                return full_name

            return full_name

        def fio_initials_surname(full_name):
            """
            Финк Оксана Анатольевна -> О.А. Финк
            Финк О.А. -> О.А. Финк
            О.А. Финк -> О.А. Финк
            """
            full_name = str(full_name or "").strip()
            parts = full_name.split()

            if len(parts) >= 3:
                return f"{parts[1][0]}.{parts[2][0]}. {parts[0]}"

            if len(parts) == 2:
                first, second = parts[0], parts[1]

                if "." in first and "." not in second:
                    return full_name

                if "." in second:
                    return f"{second} {first}"

            return full_name

        chair_fio_header = fio_surname_initials(chair_fio_raw)
        chair_fio_signature = fio_initials_surname(chair_fio_raw)
        keeper_fio_header = fio_surname_initials(keeper_fio_raw)
        keeper_fio_signature = fio_initials_surname(keeper_fio_raw)

        local_participants = self._get_participants_with_positions_from_list(
            self.temp_selected_participants
        )
        invited_participants = self._get_invited_with_positions_from_list(
            self.temp_invited_participants
        )

        def protocol_position_rank(person):
            """
            Порядок участников в протоколе после строки «Присутствовали:».

            1. Первый заместитель главы администрации города.
            2. Остальные заместители главы администрации города.
            3. Главы администраций районов.
            4. Приглашенные участники.
            5. Председатели комитетов.
            6. Начальники управлений.
            7. Заместители председателей комитетов.
            8. Начальники отделов комитетов.
            99. Остальные.
            """
            position = self._normalize_text_for_sort(person.get("dolzhnost", ""))
            source = person.get("source", "local")

            if "первый заместитель главы администрации города" in position:
                return 1

            if (
                "заместитель главы администрации города" in position
                and "первый заместитель" not in position
                and "район" not in position
            ):
                return 2

            if (
                "глава администрации" in position
                and "район" in position
                and "заместитель" not in position
            ):
                return 3

            if source == "invited":
                return 4

            if "председатель комитета" in position and "заместитель" not in position:
                return 5

            if "начальник управления" in position:
                return 6

            if "заместитель председателя" in position and "комитет" in position:
                return 7

            if "начальник отдела" in position and "комитет" in position:
                return 8

            return 99

        def protocol_sort_key(person):
            return (
                protocol_position_rank(person),
                self._person_last_name_for_sort(person),
                self._normalize_text_for_sort(person.get("fio", ""))
            )

        all_participants = sorted(local_participants + invited_participants, key=protocol_sort_key)
        participants_count = len(all_participants)

        # Актуализируем блок докладчиков/информирующих из выбранных участников.
        self.refresh_agenda_speaker_rows()
        agenda_questions = self.get_agenda_questions_for_doc()
        speakers_by_question = self._get_protocol_speakers_by_question()

        out_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"Протокол_{date_text.replace('.', '_')}.docx"
        )

        if not out_path:
            return

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "PT Astra Serif"
        style.font.size = Pt(14)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1

        section = doc.sections[0]
        # Формат листа и поля протокола:
        # справа 1 см, сверху/слева/снизу по 2 см.
        # При ширине A4 21 см рабочая область равна 18 см.
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)

        def set_font(paragraph, size=14):
            for run in paragraph.runs:
                run.font.name = "PT Astra Serif"
                run.font.size = Pt(size)
                run.bold = False

        def clear_spacing(paragraph):
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

        def add_paragraph(text="", align=None):
            p = doc.add_paragraph()
            clear_spacing(p)

            if align is not None:
                p.alignment = align

            if text:
                p.add_run(text)
                set_font(p, 14)

            return p

        def add_blank(count=1):
            for _ in range(count):
                p = doc.add_paragraph()
                clear_spacing(p)

        def set_cell_width(cell, width_cm):
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))

            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)

            tc_w.set(qn("w:w"), str(int(width_cm * 567)))
            tc_w.set(qn("w:type"), "dxa")

        def set_table_width(table, width_cm):
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            tbl_w = tbl_pr.find(qn("w:tblW"))

            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)

            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(int(width_cm * 567)))

        def set_table_fixed_layout(table):
            tbl_pr = table._tbl.tblPr
            tbl_layout = tbl_pr.find(qn("w:tblLayout"))

            if tbl_layout is None:
                tbl_layout = OxmlElement("w:tblLayout")
                tbl_pr.append(tbl_layout)

            tbl_layout.set(qn("w:type"), "fixed")

        def remove_table_borders(table):
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._tc.get_or_add_tcPr()

                    old_borders = tc_pr.find(qn("w:tcBorders"))
                    if old_borders is not None:
                        tc_pr.remove(old_borders)

                    borders = OxmlElement("w:tcBorders")

                    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "nil")
                        borders.append(border)

                    tc_pr.append(borders)

        def apply_cell_format(cell, align=WD_ALIGN_PARAGRAPH.LEFT):
            for paragraph in cell.paragraphs:
                clear_spacing(paragraph)
                paragraph.alignment = align
                set_font(paragraph, 14)

        def add_date_number_line():
            p = add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
            p.add_run(date_text)
            p.add_run(f"\t№{protocol_number}")
            set_font(p, 14)

        def add_participants_table(people):
            """
            Таблица присутствующих без видимых границ.

            Один участник = одна строка таблицы. Для каждой строки включено
            w:cantSplit, поэтому ФИО и должность одного человека не разрываются
            между двумя страницами: если строка не помещается, Word переносит ее
            целиком на следующую страницу.
            """
            if not people:
                return

            def keep_row_together(row):
                tr_pr = row._tr.get_or_add_trPr()

                old_cant_split = tr_pr.find(qn("w:cantSplit"))
                if old_cant_split is None:
                    cant_split = OxmlElement("w:cantSplit")
                    tr_pr.append(cant_split)

            def keep_paragraph_together(paragraph):
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = False

            table = doc.add_table(rows=0, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            table.allow_autofit = False
            set_table_fixed_layout(table)
            # Ширина рабочей области листа: 21 см - 2 см слева - 1 см справа = 18 см.
            # Поэтому таблица идет до правого края документа.
            set_table_width(table, 18.0)
            remove_table_borders(table)

            for person in people:
                surname, name_patronymic = self._split_fio_for_protocol_table(person.get("fio", ""))
                position = str(person.get("dolzhnost", "")).strip()

                row = table.add_row()
                keep_row_together(row)
                set_cell_width(row.cells[0], 6.8)
                set_cell_width(row.cells[1], 11.2)

                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

                left_cell = row.cells[0]
                right_cell = row.cells[1]

                left_cell.text = ""
                p_surname = left_cell.paragraphs[0]
                clear_spacing(p_surname)
                keep_paragraph_together(p_surname)
                p_surname.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_surname.add_run(surname)
                set_font(p_surname, 14)

                p_name = left_cell.add_paragraph()
                clear_spacing(p_name)
                keep_paragraph_together(p_name)
                p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_name.add_run(name_patronymic)
                set_font(p_name, 14)

                # Пустая строка после имени-отчества — это Enter между участниками.
                p_empty = left_cell.add_paragraph()
                clear_spacing(p_empty)
                keep_paragraph_together(p_empty)
                p_empty.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_font(p_empty, 14)

                right_cell.text = position
                apply_cell_format(right_cell, WD_ALIGN_PARAGRAPH.LEFT)

                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        keep_paragraph_together(paragraph)

            remove_table_borders(table)

        def add_topic_block():
            """
            Блок "совещания по вопросу ..." оформляется узким блоком
            у левого края, с выравниванием по ширине, как на вашем примере.
            """
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            table.allow_autofit = False
            set_table_fixed_layout(table)
            set_table_width(table, 8.2)
            remove_table_borders(table)

            cell = table.rows[0].cells[0]
            set_cell_width(cell, 8.2)
            cell.text = ""

            p = cell.paragraphs[0]
            clear_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(f"совещания по вопросу {title_text}")
            set_font(p, 14)

            remove_table_borders(table)

        def add_signature_line(label, value):
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            table.allow_autofit = False
            set_table_fixed_layout(table)

            # Рабочая область A4 при полях 2 см слева и 1 см справа: 18 см.
            # Поэтому правая ячейка заканчивается ровно у правого поля документа,
            # то есть фамилия будет на расстоянии 1 см от края листа.
            set_table_width(table, 18.0)

            def set_cell_margins_zero(cell):
                tc_pr = cell._tc.get_or_add_tcPr()

                old_mar = tc_pr.find(qn("w:tcMar"))
                if old_mar is not None:
                    tc_pr.remove(old_mar)

                tc_mar = OxmlElement("w:tcMar")
                for margin_name in ["top", "left", "bottom", "right"]:
                    margin = OxmlElement(f"w:{margin_name}")
                    margin.set(qn("w:w"), "0")
                    margin.set(qn("w:type"), "dxa")
                    tc_mar.append(margin)

                tc_pr.append(tc_mar)

            remove_table_borders(table)

            left = table.rows[0].cells[0]
            right = table.rows[0].cells[1]
            set_cell_width(left, 11.0)
            set_cell_width(right, 7.0)
            set_cell_margins_zero(left)
            set_cell_margins_zero(right)

            left.text = label
            right.text = value

            apply_cell_format(left, WD_ALIGN_PARAGRAPH.LEFT)
            apply_cell_format(right, WD_ALIGN_PARAGRAPH.RIGHT)

        # ===== ШАПКА =====
        add_paragraph("АДМИНИСТРАЦИЯ ГОРОДА БАРНАУЛА", WD_ALIGN_PARAGRAPH.CENTER)
        add_blank(1)
        add_paragraph("ПРОТОКОЛ", WD_ALIGN_PARAGRAPH.CENTER)
        add_blank(1)
        add_date_number_line()
        add_blank(1)

        add_topic_block()

        add_blank(1)

        add_paragraph(
            f"Председательствующий - {chair_fio_header}, {chair_position}",
            WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        add_paragraph(
            f"Протокол вела - {keeper_fio_header}, главный специалист отдела судебной работы правового комитета",
            WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        add_blank(1)

        add_paragraph("Присутствовали:", WD_ALIGN_PARAGRAPH.LEFT)

        # Участники протокола всегда выводятся списком в заданной
        # последовательности должностей, даже если их 15 и более.
        add_participants_table(all_participants)

        add_blank(1)

        add_paragraph("ПОВЕСТКА ДНЯ:", WD_ALIGN_PARAGRAPH.LEFT)

        for index, question_text in enumerate(agenda_questions, start=1):
            add_paragraph(f"{index}.\t{question_text}", WD_ALIGN_PARAGRAPH.JUSTIFY)

            speakers = speakers_by_question.get(index, [])

            if speakers:
                label = "Докладывает:" if len(speakers) == 1 else "Докладывают:"
                p = add_paragraph(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                p.add_run(label + " ")

                for speaker_index, speaker in enumerate(speakers):
                    text = f"{speaker.get('dolzhnost', '')} - {self._short_fio(speaker.get('fio', ''))}"
                    if speaker_index == 0:
                        p.add_run(text)
                    else:
                        p.add_run(", " + text)

                set_font(p, 14)

        add_blank(1)
        add_paragraph("Заслушав информацию,", WD_ALIGN_PARAGRAPH.JUSTIFY)
        add_blank(1)

        add_paragraph("1.\tРЕШИЛИ:", WD_ALIGN_PARAGRAPH.JUSTIFY)
        add_paragraph("1.1.\tИнформацию принять к сведению.", WD_ALIGN_PARAGRAPH.JUSTIFY)

        # В документ протокола переносим поручения из блока "Поручения".
        # Если выбран один или несколько комитетов, пункт начинается с:
        # "Комитету ... совместно с комитетом ... в срок до ...".
        protocol_tasks = []

        for task_row in getattr(self, "tasks_rows", []):
            try:
                task_text = task_row["task"].get().strip()
            except Exception:
                task_text = ""

            try:
                date_text_task = task_row.get("date").get().strip() if task_row.get("date") else ""
            except Exception:
                date_text_task = ""

            committees = self._get_task_committees_from_row(task_row)

            if not task_text and not committees:
                continue

            committee_prefix = self._format_protocol_committee_assignment(committees)

            if committee_prefix:
                full_text = committee_prefix

                if date_text_task:
                    full_text += f" в срок до {date_text_task}"

                if task_text:
                    full_text += f" {task_text}"

                if full_text and not full_text.endswith((".", "!", "?")):
                    full_text += "."

                protocol_tasks.append(full_text)
            elif task_text:
                full_text = task_text
                if full_text and not full_text.endswith((".", "!", "?")):
                    full_text += "."
                protocol_tasks.append(full_text)

        # Дополнительное итоговое поручение под блоком "Поручения".
        # Оно добавляется последним пунктом, если стоит галочка.
        try:
            report_enabled = self.protocol_report_enabled.get()
        except Exception:
            report_enabled = False

        if report_enabled:
            report_text = self.protocol_report_text.get().strip()
            report_date = self.protocol_report_date.get().strip()

            if report_text:
                final_report = report_text
                if report_date:
                    final_report += report_date
                if final_report and not final_report.endswith((".", "!", "?")):
                    final_report += "."
                protocol_tasks.append(final_report)

        if protocol_tasks:
            for task_index, task_text in enumerate(protocol_tasks, start=2):
                p_task = add_paragraph(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                p_task.add_run(f"1.{task_index}.\t{task_text}")
                set_font(p_task, 14)
        else:
            add_paragraph("1.2.\t", WD_ALIGN_PARAGRAPH.JUSTIFY)

        add_blank(2)

        add_signature_line("Председательствующий", chair_fio_signature)
        add_blank(2)
        add_signature_line("Протокол вела", keeper_fio_signature)

        saved_path = self.safe_save_docx(doc, out_path)

        if not saved_path:
            return

        messagebox.showinfo("Успех", f"✅ Протокол создан:\n{saved_path}")

    def open_protocol_window(self):
        """
        Совместимость со старой кнопкой: теперь протокол открывается
        не отдельным окном, а как раздел в основном окне.
        """
        self.attachment_type.set("protocol")

        for doc_type, btn in self.doc_type_buttons.items():
            btn.configure(style="Selected.TButton" if doc_type == "protocol" else "Modern.TButton")

        self.update_visible_fields_by_doc_type()

    def update_visible_fields_by_doc_type(self):
        """
        Переключает внешний вид основного окна под выбранный тип документа.
        Кнопки выбора документа всегда находятся в верхней строке.
        """
        selected = self.attachment_type.get()

        # Скрываем контейнер рабочих блоков.
        # Для разделов "Список участников" и "Телефонограмма" он не нужен;
        # иначе пустой row2 оставляет лишний зазор перед таблицей "Архив совещаний".
        try:
            self.row2.pack_forget()
        except Exception:
            pass

        # Сначала скрываем все переключаемые блоки.
        for widget_name in [
            "left_col",
            "right_col",
            "agenda_label",
            "agenda_block",
            "protocol_label",
            "protocol_block",
            "email_block",
            "transfer_block",
            "agenda_sign_block",
            "protocol_keeper_block",
            "protocol_inline_block",
            "participants_select_frame",
            "document_bar",
            "top_telegram_number_block",
            "top_protocol_number_block",
            "protocol_report_frame",
        ]:
            widget = getattr(self, widget_name, None)
            if widget:
                try:
                    widget.pack_forget()
                except Exception:
                    pass

        if hasattr(self, "agenda_add_btn"):
            try:
                self.agenda_add_btn.master.pack_forget()
            except Exception:
                pass

        if hasattr(self, "task_btns"):
            try:
                self.task_btns.pack_forget()
            except Exception:
                pass

        # По умолчанию поле "Протокол вела" на строке Дата/Время/Кабинет скрыто.
        for protocol_keeper_widget in (
            getattr(self, "row_date_protocol_keeper_label", None),
            getattr(self, "row_date_protocol_keeper_entry", None),
        ):
            if protocol_keeper_widget:
                try:
                    protocol_keeper_widget.pack_forget()
                except Exception:
                    pass

        # Текущая точка вставки внутри form_container.
        # row_date всегда остается упакованной, поэтому используем after=...,
        # а не before=self.row2: row2 в некоторых разделах скрыт, и Tkinter
        # выдает ошибку "isn't packed" при попытке вставить блок перед ним.
        current_after = self.row_date

        def pack_top(widget, pady=(0, 4), fill="x", expand=False):
            nonlocal current_after
            if not widget:
                return
            widget.pack(fill=fill, expand=expand, pady=pady, after=current_after)
            current_after = widget

        def pack_document_controls_above_archive():
            # Во всех разделах панель документов должна идти сразу перед архивом
            # совещаний. Если row2 скрыт, она будет сразу после служебных полей.
            # Если row2 показан, она будет сразу после рабочих блоков.
            pack_top(self.document_bar, pady=(2, 4))

        # =====================================================
        # СПИСОК УЧАСТНИКОВ
        # =====================================================
        if selected == "participants":
            # "Протокол ведет" — прямо под полями "По вопросу / Дата / Время / Кабинет".
            pack_top(self.protocol_keeper_block)

            # Кнопки выбора участников — ниже поля "Протокол ведет".
            pack_top(self.participants_select_frame)

            # Блок документов должен идти сразу под кнопками
            # "Органы МСУ" / "Приглашенные", а архив — сразу под ним.
            pack_document_controls_above_archive()

        # =====================================================
        # ПОВЕСТКА
        # =====================================================
        elif selected == "agenda":
            # "Должность" и "ФИО" — прямо под верхними полями.
            pack_top(self.agenda_sign_block)

            # Рабочие блоки повестки.
            pack_top(self.row2, pady=(2, 5), fill="both", expand=True)
            self.left_col.pack(side="left", fill="both", expand=True, padx=(0, 5))
            self.right_col.pack(side="right", fill="both", expand=True, padx=(5, 0))

            self.agenda_label.config(text="Вопрос в повестке:")
            self.agenda_label.pack(anchor="w")

            self.agenda_block.pack(fill="x", expand=False, pady=(2, 0))
            self.agenda_block.pack_propagate(False)

            if hasattr(self, "agenda_add_btn"):
                self.agenda_add_btn.master.pack(fill="x")

            self.tasks_label.config(text="Выбранные участники совещания:")
            self.tasks_label.pack(anchor="w", padx=5)
            self.tasks_block.pack(fill="x", expand=False, pady=(2, 0))
            self.tasks_block.pack_propagate(False)

            self.refresh_agenda_speaker_rows()

            if hasattr(self, "render_agenda_speaker_rows"):
                self.render_agenda_speaker_rows()

            pack_document_controls_above_archive()

        # =====================================================
        # ТЕЛЕФОНОГРАММА
        # =====================================================
        elif selected == "telegram":
            # Номер телефонограммы — над полем "По вопросу".
            self.top_telegram_number_block.pack(fill="x", pady=(0, 4), before=self.row_title)

            # Электронная почта / Передала / Тел. / ФИО — прямо под верхними полями.
            pack_top(self.transfer_block)

            # Блок документов должен быть поднят сразу под полями
            # телефонограммы, а архив — сразу под ним.
            pack_document_controls_above_archive()

        # =====================================================
        # ПРОТОКОЛ
        # =====================================================
        elif selected == "protocol":
            # Номер протокола — над полем "По вопросу".
            self.top_protocol_number_block.pack(fill="x", pady=(0, 4), before=self.row_title)

            # Поле "Протокол вела" — на строке Дата / Время / Кабинет, справа от "Кабинет".
            self.row_date_protocol_keeper_label.pack(side="left", padx=(12, 0))
            self.row_date_protocol_keeper_entry.pack(side="left", padx=5)

            # Председательствующий / Должность — прямо под верхними полями.
            pack_top(self.protocol_inline_block)

            # Блок "Поручения" ниже служебных полей.
            pack_top(self.row2, pady=(2, 5), fill="both", expand=True)

            # Блок "Поручения".
            self.right_col.pack(side="left", fill="both", expand=True, padx=0)
            self.tasks_label.config(text="Поручения:")
            self.tasks_label.pack(anchor="w", padx=5)
            self.tasks_block.pack(fill="x", expand=False, pady=(2, 0))
            self.tasks_block.pack_propagate(False)

            if hasattr(self, "task_btns"):
                self.task_btns.pack(fill="x")

            if hasattr(self, "protocol_report_frame"):
                self.protocol_report_frame.pack(fill="x", pady=(2, 0))
                if hasattr(self, "update_protocol_report_state"):
                    self.update_protocol_report_state()

            if hasattr(self, "render_task_rows"):
                self.render_task_rows()

            pack_document_controls_above_archive()

        # =====================================================
        # НИЧЕГО НЕ ВЫБРАНО
        # =====================================================
        else:
            pack_top(self.row2, pady=(2, 5), fill="both", expand=True)
            self.right_col.pack(side="left", fill="both", expand=True, padx=0)
            self.tasks_label.config(text="Поручения:")
            self.tasks_label.pack(anchor="w", padx=5)
            self.tasks_block.pack(fill="x", expand=False, pady=(2, 0))
            self.tasks_block.pack_propagate(False)

            if hasattr(self, "task_btns"):
                self.task_btns.pack(fill="x")

            if hasattr(self, "render_task_rows"):
                self.render_task_rows()

            self.participants_select_frame.pack(fill="x", pady=(0, 4))
            pack_document_controls_above_archive()

    def generate_document(self):
        t = self.attachment_type.get()

        if t == "participants":
            self.generate_participants_doc()

        elif t == "agenda":
            self.generate_agenda_doc()


        elif t == "telegram":
            self.generate_telegram_doc()

        elif t == "protocol":
            self.generate_protocol_doc()

        else:
            messagebox.showwarning("Выбор", "Выберите тип документа")

    def open_user_selector(self):
        current = "\n".join(self.temp_selected_participants) if self.temp_selected_participants else ""
        SelectParticipantsWindow(self, current, target="local")

    def open_invited_selector(self):
        current = "\n".join(self.temp_invited_participants) if self.temp_invited_participants else ""
        SelectParticipantsWindow(self, current, target="invited")

    def on_participants_selected(self, selected_list):
        self.temp_selected_participants = selected_list

        if self.attachment_type.get() == "agenda":
            self.refresh_agenda_speaker_rows()

        messagebox.showinfo(
            "Успех",
            f"Выбрано сотрудников для блока 'Органы местного самоуправления': {len(selected_list)}"
        )

    def on_invited_selected(self, selected_list):
        self.temp_invited_participants = selected_list

        if self.attachment_type.get() == "agenda":
            self.refresh_agenda_speaker_rows()

        messagebox.showinfo(
            "Успех",
            f"Выбрано сотрудников для блока 'Приглашенные': {len(selected_list)}"
        )

    def _open_file_with_default_app(self, file_path):
        """
        Открывает файл в той программе, которая назначена для его формата
        в системе Windows / macOS / Linux.
        """
        import os
        import sys
        import subprocess

        file_path = os.path.abspath(file_path)

        if not os.path.exists(file_path):
            messagebox.showerror("Ошибка", f"Файл не найден:\n{file_path}")
            return

        try:
            if sys.platform.startswith("win"):
                os.startfile(file_path)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", file_path])
            else:
                subprocess.Popen(["xdg-open", file_path])

        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Не удалось открыть файл:\n{file_path}\n\n{e}"
            )

    def open_attached_file(self):
        """
        Открывает прикрепленный документ.
        Если прикреплен один файл — открывает сразу.
        Если файлов несколько — показывает список для выбора.
        """
        files = []

        for path in getattr(self, "attached_files", []):
            path = str(path or "").strip()
            if path:
                files.append(path)

        if not files and getattr(self, "current_file_path", ""):
            files = [
                p.strip()
                for p in str(self.current_file_path).split("\n")
                if p.strip()
            ]

        if not files:
            messagebox.showinfo("Документы", "Нет прикрепленных документов.")
            return

        existing_files = [p for p in files if os.path.exists(p)]

        if not existing_files:
            messagebox.showerror(
                "Ошибка",
                "Прикрепленные файлы не найдены.\n\n"
                "Возможно, они были перемещены или удалены."
            )
            return

        if len(existing_files) == 1:
            self._open_file_with_default_app(existing_files[0])
            return

        select_win = tk.Toplevel(self)
        select_win.title("Открыть документ")
        select_win.geometry("700x360")
        select_win.transient(self)
        select_win.grab_set()
        select_win.configure(bg="#f4f6f9")

        ttk.Label(
            select_win,
            text="Выберите документ для открытия:",
            font=("Segoe UI", 11, "bold")
        ).pack(anchor="w", padx=10, pady=(10, 5))

        list_frame = tk.Frame(select_win, bg="#f4f6f9")
        list_frame.pack(fill="both", expand=True, padx=10, pady=5)

        listbox = tk.Listbox(
            list_frame,
            font=("Segoe UI", 10),
            activestyle="dotbox"
        )
        listbox.pack(side="left", fill="both", expand=True)

        scroll = ttk.Scrollbar(list_frame, orient="vertical", command=listbox.yview)
        scroll.pack(side="right", fill="y")
        listbox.configure(yscrollcommand=scroll.set)

        for file_path in existing_files:
            listbox.insert("end", os.path.basename(file_path))

        def open_selected(event=None):
            selection = listbox.curselection()
            if not selection:
                messagebox.showinfo("Документы", "Выберите документ из списка.")
                return

            file_path = existing_files[selection[0]]
            select_win.destroy()
            self._open_file_with_default_app(file_path)

        btn_frame = tk.Frame(select_win, bg="#f4f6f9")
        btn_frame.pack(fill="x", padx=10, pady=10)

        ttk.Button(
            btn_frame,
            text="📂 Открыть",
            command=open_selected,
            style="Modern.TButton"
        ).pack(side="right", padx=5)

        ttk.Button(
            btn_frame,
            text="❌ Отмена",
            command=select_win.destroy,
            style="Modern.TButton"
        ).pack(side="right", padx=5)

        listbox.bind("<Double-1>", open_selected)

        if existing_files:
            listbox.selection_set(0)
            listbox.focus_set()

    def attach_file(self):
        date_val = self.ent_date.get().strip()

        if not date_val:
            messagebox.showwarning("Внимание", "Сначала укажите дату совещания.")
            return

        files = filedialog.askopenfilenames(title="Выберите файлы")

        if not files:
            return

        folder_name = date_val.replace("/", ".").replace("\\", ".").replace(":", ".")
        target_dir = os.path.join("documents", folder_name)

        os.makedirs(target_dir, exist_ok=True)

        copied_files = []

        for file_path in files:
            try:
                file_name = os.path.basename(file_path)
                target_path = os.path.join(target_dir, file_name)

                if os.path.exists(target_path):
                    name, ext = os.path.splitext(file_name)
                    counter = 1

                    while os.path.exists(target_path):
                        new_name = f"{name}_{counter}{ext}"
                        target_path = os.path.join(target_dir, new_name)
                        counter += 1

                shutil.copy2(file_path, target_path)
                copied_files.append(target_path)

            except Exception as e:
                messagebox.showerror(
                    "Ошибка",
                    f"Не удалось прикрепить файл:\n{file_path}\n\n{e}"
                )

        if copied_files:
            self.attached_files.extend(copied_files)
            self.current_file_path = "\n".join(self.attached_files)

            names = [os.path.basename(p) for p in self.attached_files]
            preview = "\n".join(names[:3])

            if len(names) > 3:
                preview += f"\nи ещё {len(names) - 3}"

            self.attachment_var.set(preview)

            messagebox.showinfo(
                "Успех",
                f"Файлы прикреплены и скопированы в папку:\n{target_dir}"
            )

    def load_meetings(self):
        for i in self.tree.get_children():
            self.tree.delete(i)

        def meeting_status_symbol(protocol_text):
            has_dated_task = False
            has_unchecked_dated_task = False

            for line in str(protocol_text or "").split("\n"):
                clean = line.strip()

                if not clean:
                    continue

                if ". " in clean:
                    clean = clean.split(". ", 1)[1]

                parts = [part.strip() for part in clean.split("|")]

                if len(parts) >= 5:
                    date_text = parts[3]
                    done_text = parts[4]
                elif len(parts) >= 4:
                    date_text = parts[2]
                    done_text = parts[3]
                elif len(parts) >= 3:
                    date_text = parts[1]
                    done_text = parts[2]
                else:
                    continue

                if date_text:
                    has_dated_task = True

                    if "☑" not in done_text:
                        has_unchecked_dated_task = True

            if has_unchecked_dated_task:
                return "!"

            if has_dated_task:
                return "✓"

            return ""

        try:
            self.tree.tag_configure("task_warning", foreground="#c00000")
            self.tree.tag_configure("task_done", foreground="#148a2b")

            for idx, row in enumerate(meetings_get_all(), 1):
                parts = [p.strip() for p in (row.get('participants', '') or '').split('\n') if p.strip()]
                preview = ", ".join(parts[:2]) + (f" и ещё {len(parts) - 2}" if len(parts) > 2 else "")

                status_symbol = meeting_status_symbol(row.get('protocol', '') or '')
                display_idx = f"{status_symbol} {idx}" if status_symbol else str(idx)

                row_tags = [row['id']]

                if status_symbol == "!":
                    row_tags.append("task_warning")
                elif status_symbol == "✓":
                    row_tags.append("task_done")

                self.tree.insert(
                    "",
                    "end",
                    values=(display_idx, row['date_val'], row['time_val'], row['title'], preview),
                    tags=tuple(row_tags)
                )
        except Exception as e:
            messagebox.showerror("Ошибка", f"Загрузка списка: {e}")

    def on_select(self, event):
        item = self.tree.focus()
        if item:
            tags = self.tree.item(item, "tags")
            if tags:
                self.load_meeting_details(int(tags[0]))

    def load_meeting_details(self, mid):
        try:
            row = next((m for m in meetings_get_all() if m['id'] == mid), None)
            if not row:
                return

            self.current_id = mid

            self.ent_title.delete(0, "end")
            self.ent_title.insert(0, row['title'])

            self.ent_date.delete(0, "end")
            self.ent_date.insert(0, row['date_val'])

            self.ent_time.delete(0, "end")
            self.ent_time.insert(0, row['time_val'])

            self.protocol_keeper.set(
                row.get("protocol_keeper", "") or "Иванова Елена Николаевна"
            )

            self._set_agenda_questions_from_text(row.get('agenda', '') or "")

            # ===== ЗАГРУЗКА ПОРУЧЕНИЙ =====
            self.tasks_rows.clear()
            self.protocol_report_enabled.set(False)
            self.protocol_report_text.set(
                "О проделанной работе проинформировать правовой комитет администрации города Барнаула до "
            )
            self.protocol_report_date.set("")
            if hasattr(self, "update_protocol_report_state"):
                self.update_protocol_report_state()

            protocol_text = row.get('protocol', '')
            lines = [line.strip() for line in protocol_text.split("\n") if line.strip()]

            if lines:
                for line in lines:
                    clean = line

                    if clean.startswith("__REPORT__"):
                        parts = [p.strip() for p in clean.split("|")]
                        self.protocol_report_enabled.set(True)
                        if len(parts) > 1 and parts[1]:
                            self.protocol_report_text.set(parts[1])
                        if len(parts) > 2:
                            self.protocol_report_date.set(parts[2])
                        if hasattr(self, "update_protocol_report_state"):
                            self.update_protocol_report_state()
                        continue

                    if ". " in clean:
                        clean = clean.split(". ", 1)[1]

                    parts = [p.strip() for p in clean.split("|")]

                    committee_text = ""
                    extra_committee_values = []
                    task_text = ""
                    date_text = ""
                    done = False

                    if len(parts) >= 5:
                        # Старый формат: комитет | второй комитет | поручение | дата | галочка
                        committee_text = parts[0]
                        if parts[1]:
                            extra_committee_values = [parts[1]]
                        task_text = parts[2]
                        date_text = parts[3]
                        done = "☑" in parts[4]
                    elif len(parts) >= 4:
                        # Новый формат: комитеты через §§ | поручение | дата | галочка
                        committees = [value.strip() for value in parts[0].split("§§") if value.strip()]
                        committee_text = committees[0] if committees else ""
                        extra_committee_values = committees[1:]
                        task_text = parts[1]
                        date_text = parts[2]
                        done = "☑" in parts[3]
                    else:
                        committee_text = ""
                        task_text = parts[0] if len(parts) > 0 else ""
                        date_text = parts[1] if len(parts) > 1 else ""
                        done = "☑" in parts[2] if len(parts) > 2 else False

                    self.add_task_row(task_text, date_text, done, committee_text, extra_committee_values)
            else:
                self.add_task_row()

            self.attachment_type.set("none")

            if row.get('has_participants_list'):
                self.attachment_type.set("participants")
            elif row.get('has_telegram'):
                self.attachment_type.set("telegram")


            self.update_visible_fields_by_doc_type()

            for doc_type, btn in self.doc_type_buttons.items():
                if doc_type == self.attachment_type.get():
                    btn.configure(style="Selected.TButton")
                else:
                    btn.configure(style="Modern.TButton")

            self.temp_selected_participants = [
                p.strip()
                for p in (row.get('participants', '') or '').split('\n')
                if p.strip()
            ]
            self.temp_invited_participants = [
                p.strip()
                for p in (row.get("invited_participants", "") or "").split("\n")
                if p.strip()
            ]

            self.current_file_path = row.get('attachment_path', '') or ""

            self.attached_files = [
                p.strip()
                for p in self.current_file_path.split("\n")
                if p.strip()
            ]

            if self.attached_files:
                names = [os.path.basename(p) for p in self.attached_files]
                preview = "\n".join(names[:3])

                if len(names) > 3:
                    preview += f"\nи ещё {len(names) - 3}"

                self.attachment_var.set(preview)
            else:
                self.attachment_var.set("Файл не прикреплен")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Загрузка деталей: {e}")

    def _get_participants_with_positions(self):
        try:
            return self._get_participants_with_positions_from_list(
                self.temp_selected_participants
            )
        except Exception as e:
            messagebox.showerror("БД", f"Загрузка должностей: {e}")
            return []

    def _apply_no_line_end_prepositions(self, doc):
        """
        Во всех формируемых Word-документах не допускает ситуацию,
        когда строка заканчивается коротким предлогом или союзом.

        Технически это делается через неразрывный пробел после таких слов:
        "в работу" -> "в\u00A0работу". Word не сможет оставить "в" в конце
        строки и перенесет вместе со следующим словом.
        """
        import re

        # Предлоги, союзы и близкие служебные слова, которые не должны
        # оставаться последними в строке.
        service_words = [
            "а", "без", "близ", "в", "во", "вместо", "вне", "для", "до",
            "за", "и", "или", "из", "из-за", "из-под", "к", "ко", "как",
            "ли", "либо", "меж", "между", "на", "над", "надо", "не", "ни",
            "но", "о", "об", "обо", "от", "ото", "перед", "передо", "по",
            "под", "подо", "при", "про", "с", "со", "у", "через", "что",
            "чтобы", "же", "то"
        ]

        escaped = sorted((re.escape(w) for w in service_words), key=len, reverse=True)
        # Ищем служебное слово как отдельное слово и обычный пробел после него.
        # Пробел заменяем на неразрывный.
        pattern = re.compile(
            r"(?i)(?<![A-Za-zА-Яа-яЁё0-9_])(" + "|".join(escaped) + r")(?![A-Za-zА-Яа-яЁё0-9_])\s+"
        )

        def fix_text(text):
            if not text or " " not in text:
                return text

            text = pattern.sub(lambda m: m.group(1) + "\u00A0", text)

            # Сокращения тоже не должны оставаться в конце строки.
            # После них ставим неразрывный пробел.
            abbreviations = [
                "ул.", "каб.", "час.", "г.", "д.", "стр.", "п.",
                "т.п.", "т.д.", "и т.п.", "и т.д."
            ]
            for abbr in abbreviations:
                text = text.replace(abbr + " ", abbr + "\u00A0")

            # Инициалы и фамилия одного человека не должны разделяться переносом строки:
            # "С.В. Рябчуну" -> "С.В.\u00A0Рябчуну".
            text = re.sub(
                r"([А-ЯЁ]\.[А-ЯЁ]\.)\s+([А-ЯЁ][а-яё]+)",
                lambda m: m.group(1) + "\u00A0" + m.group(2),
                text
            )

            # Фамилия и инициалы тоже держим вместе: "Иванов И.И.".
            text = re.sub(
                r"([А-ЯЁ][а-яё]+)\s+([А-ЯЁ]\.[А-ЯЁ]\.)",
                lambda m: m.group(1) + "\u00A0" + m.group(2),
                text
            )

            # Дата не должна оставаться в конце строки отдельно от предыдущего слова:
            # "состоится 21.05.2026" -> "состоится\u00A021.05.2026".
            text = re.sub(
                r"\s+(\d{1,2}\.\d{1,2}\.\d{4})(?=\D|$)",
                lambda m: "\u00A0" + m.group(1),
                text
            )

            # Короткие числовые даты с годом через пробел также держим вместе
            # с предыдущим словом: "20 мая 2026".
            text = re.sub(
                r"\s+(\d{1,2}\s+(?:января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4})(?=\D|$)",
                lambda m: "\u00A0" + m.group(1).replace(" ", "\u00A0"),
                text,
                flags=re.IGNORECASE
            )

            return text

        def fix_paragraph(paragraph):
            for run in paragraph.runs:
                new_text = fix_text(run.text)
                if new_text != run.text:
                    run.text = new_text

        def fix_table(table):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        fix_paragraph(paragraph)
                    for inner_table in cell.tables:
                        fix_table(inner_table)

        for paragraph in doc.paragraphs:
            fix_paragraph(paragraph)

        for table in doc.tables:
            fix_table(table)

    def safe_save_docx(self, doc, out_path):
        """
        Безопасно сохраняет docx.
        Если файл уже открыт в Word или нет доступа к пути,
        предлагает выбрать другое имя файла.
        """
        from tkinter import filedialog, messagebox
        import os

        while True:
            try:
                self._apply_no_line_end_prepositions(doc)
                doc.save(out_path)
                return out_path

            except PermissionError:
                messagebox.showwarning(
                    "Файл занят",
                    "Не удалось сохранить документ.\n\n"
                    "Скорее всего, файл уже открыт в Word или занят другой программой.\n"
                    "Закройте файл или выберите другое имя для сохранения."
                )

                folder = os.path.dirname(out_path)
                filename = os.path.basename(out_path)
                name, ext = os.path.splitext(filename)

                new_path = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word Document", "*.docx")],
                    initialdir=folder if folder else None,
                    initialfile=f"{name}_новый{ext}"
                )

                if not new_path:
                    return None

                out_path = new_path

            except Exception as e:
                messagebox.showerror(
                    "Ошибка сохранения",
                    f"Не удалось сохранить документ:\n{e}"
                )
                return None

    def _get_employee_by_name(self, full_name):
        try:
            all_emps = get_all_employees()
            return next(
                (e for e in all_emps if e.get("full_name", "").strip() == full_name.strip()),
                None
            )
        except Exception:
            return None

    def _normalize_text_for_sort(self, text):
        """
        Нормализует текст для поиска должности:
        - нижний регистр;
        - ё -> е;
        - убирает лишние пробелы.
        """
        text = str(text or "").strip().lower()
        text = text.replace("ё", "е")
        text = " ".join(text.split())
        return text

    def _person_last_name_for_sort(self, person):
        """
        Возвращает фамилию для сортировки по алфавиту.
        Если ФИО неполное, сортирует по всей строке.
        """
        fio = self._normalize_text_for_sort(person.get("fio", ""))
        parts = fio.split()

        if parts:
            return parts[0]

        return fio

    def _official_position_rank(self, position):
        """
        Единый порядок сортировки должностей.

        1. Заместитель главы
        2. Глава администрации района
        3. Заместитель главы администрации района
        4. Председатель комитета
        5. Заместитель председателя комитета
        6. Начальник управления
        7. Заведующий отделом
        8. Начальник отдела
        99. Все остальные
        """
        p = self._normalize_text_for_sort(position)

        # 1. Заместители главы администрации города:
        # первый заместитель главы администрации города,
        # заместитель главы администрации города,
        # заместитель главы администрации города, руководитель аппарата,
        # заместитель главы администрации города по ...
        if (
                "первый заместитель главы администрации города" in p
                or "заместитель главы администрации города" in p
                or (
                "заместитель главы" in p
                and "администрации района" not in p
        )
        ):
            return 1

        # 2. Главы администраций районов
        if (
                "глава администрации района" in p
                or "глава администрации индустриального района" in p
                or "глава администрации октябрьского района" in p
                or "глава администрации железнодорожного района" in p
                or "глава администрации ленинского района" in p
                or "глава администрации центрального района" in p
        ):
            return 2

        # 3. Заместители глав администраций районов
        if (
                "заместитель главы администрации района" in p
                or "заместитель главы администрации индустриального района" in p
                or "заместитель главы администрации октябрьского района" in p
                or "заместитель главы администрации железнодорожного района" in p
                or "заместитель главы администрации ленинского района" in p
                or "заместитель главы администрации центрального района" in p
        ):
            return 3

        # Важно: заместитель председателя должен проверяться ДО председателя
        if "заместитель председателя" in p and "комитет" in p:
            return 5

        if (
                "председатель комитета" in p
                or p.startswith("председатель ")
                and "комитет" in p
        ):
            return 4

        if "начальник управления" in p:
            return 6

        if "заведующий отделом" in p:
            return 7

        if "начальник отдела" in p:
            return 8

        return 99

    def _official_position_group_title(self, position):
        """
        Название подблока в документе.
        """
        rank = self._official_position_rank(position)

        titles = {
            1:"Заместители главы:",
            2:"Главы администраций районов:",
            3:"Заместители глав администраций районов:",
            4:"Председатели комитетов:",
            5:"Заместители председателей комитетов:",
            6:"Начальники управлений:",
            7:"Заведующие отделами:",
            8:"Начальники отделов:",
            99:"Иные участники:"
        }

        return titles.get(rank, "Иные участники:")

    def _sort_people_by_official_rank(self, participants):
        """
        Единая сортировка для:
        - списка участников;
        - блока докладывающих / информирующих в повестке;
        - телефонограммы.

        Сначала по рангу должности, затем по фамилии.
        """
        return sorted(
            participants,
            key=lambda p:(
                self._official_position_rank(p.get("dolzhnost", "")),
                self._person_last_name_for_sort(p),
                self._normalize_text_for_sort(p.get("fio", ""))
            )
        )

    def _group_people_by_official_rank(self, participants):
        """
        Возвращает участников группами по должности.
        Внутри каждой группы участники уже отсортированы по алфавиту.
        """
        sorted_people = self._sort_people_by_official_rank(participants)

        groups = []

        for person in sorted_people:
            rank = self._official_position_rank(person.get("dolzhnost", ""))
            title = self._official_position_group_title(person.get("dolzhnost", ""))

            if not groups or groups[-1]["rank"] != rank:
                groups.append({
                    "rank":rank,
                    "title":title,
                    "people":[]
                })

            groups[-1]["people"].append(person)

        return groups

    def _sort_invited_participants(self, participants):
        """
        Приглашенные сортируются по тому же должностному порядку,
        а внутри каждого должностного блока — по алфавиту.
        """
        return self._sort_people_by_official_rank(participants)

    def _sort_local_participants(self, participants):
        """
        Органы МСУ сортируются по тому же должностному порядку,
        а внутри каждого должностного блока — по алфавиту.
        """
        return self._sort_people_by_official_rank(participants)

    def _get_invited_with_positions_from_list(self, names):
        try:
            invited_people = self._load_invited_people_records()

            if not invited_people:
                invited_people = invited_get_all()

            res = []

            for name in names:
                person = next(
                    (
                        p for p in invited_people
                        if (
                                   p.get("full_name")
                                   or p.get("fio")
                                   or p.get("name")
                                   or ""
                           ).strip() == name.strip()
                    ),
                    None
                )

                if person:
                    fio = (
                            person.get("full_name")
                            or person.get("fio")
                            or person.get("name")
                            or name
                    ).strip()

                    position = (
                            person.get("position")
                            or person.get("dolzhnost")
                            or person.get("job_title")
                            or "Не указана"
                    )

                    leadership = (
                            person.get("leadership")
                            or person.get("rukovodstvo")
                            or person.get("supervisor")
                            or person.get("boss")
                            or ""
                    ).strip()

                    res.append({
                        "fio": fio,
                        "dolzhnost": position,
                        "leadership": leadership,
                        "employee": None,
                        "source": "invited"
                    })
                else:
                    res.append({
                        "fio": name,
                        "dolzhnost": "Не найдено",
                        "leadership": "",
                        "employee": None,
                        "source": "invited"
                    })

            return res

        except Exception as e:
            messagebox.showerror("БД", f"Загрузка приглашенных: {e}")
            return []
    def _get_participants_with_positions_from_list(self, names):
        try:
            local_people = self._load_local_people_records()

            # Если в новой таблице настроек пока пусто, используем старую таблицу сотрудников.
            if not local_people:
                local_people = get_all_employees()

            res = []

            for name in names:
                emp = next(
                    (
                        e for e in local_people
                        if (
                            e.get("full_name")
                            or e.get("fio")
                            or e.get("name")
                            or ""
                        ).strip() == name.strip()
                    ),
                    None
                )

                if emp:
                    fio = (
                        emp.get("full_name")
                        or emp.get("fio")
                        or emp.get("name")
                        or name
                    ).strip()

                    position = (
                        emp.get("position")
                        or emp.get("dolzhnost")
                        or emp.get("job_title")
                        or "Не указана"
                    )

                    leadership = (
                        emp.get("leadership")
                        or emp.get("rukovodstvo")
                        or emp.get("supervisor")
                        or emp.get("boss")
                        or ""
                    ).strip()

                    res.append({
                        "fio": fio,
                        "dolzhnost": position,
                        "leadership": leadership,
                        "employee": emp,
                        "source": "local"
                    })
                else:
                    res.append({
                        "fio": name,
                        "dolzhnost": "Не найдено",
                        "leadership": "",
                        "employee": None,
                        "source": "local"
                    })

            return res

        except Exception as e:
            messagebox.showerror("БД", f"Загрузка должностей: {e}")
            return []

    def generate_participants_doc(self):
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from tkinter import filedialog, messagebox

        local_participants = self._get_participants_with_positions_from_list(
            self.temp_selected_participants
        )
        invited_participants = self._get_invited_with_positions_from_list(
            self.temp_invited_participants
        )

        local_participants = self._sort_local_participants(local_participants)
        invited_participants = self._sort_invited_participants(invited_participants)

        out_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"Список_{self.ent_date.get().replace('.', '_')}.docx"
        )

        if not out_path:
            return

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "PT Astra Serif"
        style.font.size = Pt(14)

        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)

        def set_font(paragraph, size=14, bold=False):
            for run in paragraph.runs:
                run.font.name = "PT Astra Serif"
                run.font.size = Pt(size)
                run.bold = bold

        def set_cell_font(cell, size=14):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "PT Astra Serif"
                    run.font.size = Pt(size)

        def set_cell_width(cell, width_cm):
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()

            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)

            tc_w.set(qn("w:w"), str(int(width_cm * 567)))
            tc_w.set(qn("w:type"), "dxa")

        def set_table_fixed_layout(table):
            tbl_pr = table._tbl.tblPr
            tbl_layout = tbl_pr.find(qn("w:tblLayout"))

            if tbl_layout is None:
                tbl_layout = OxmlElement("w:tblLayout")
                tbl_pr.append(tbl_layout)

            tbl_layout.set(qn("w:type"), "fixed")

        def set_table_width(table, width_cm):
            tbl = table._tbl
            tbl_pr = tbl.tblPr

            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)

            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(int(width_cm * 567)))

        def set_row_min_height(row, height_twips=520):
            """
            Делает строку таблицы выше.
            520 twips примерно 0.9 см.
            """
            tr_pr = row._tr.get_or_add_trPr()

            tr_height = tr_pr.find(qn("w:trHeight"))
            if tr_height is None:
                tr_height = OxmlElement("w:trHeight")
                tr_pr.append(tr_height)

            tr_height.set(qn("w:val"), str(height_twips))
            tr_height.set(qn("w:hRule"), "atLeast")

        def justify_cell_text(cell):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.line_spacing = 1

        def center_cell_text(cell):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1

        def format_fio_for_table(full_name):
            parts = str(full_name or "").strip().split()

            if len(parts) >= 3:
                return parts[0], " ".join(parts[1:])

            return str(full_name or "").strip(), ""

        # ===== Заголовок =====
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing = 1

        title.add_run(
            "СПИСОК\n"
            f"приглашенных на совещание по вопросу {self.ent_title.get().strip()}"
        )
        set_font(title, 14)

        info = doc.add_paragraph()
        info.paragraph_format.left_indent = Cm(10)
        info.paragraph_format.line_spacing = 1
        info.add_run(f"{self.ent_date.get()}\n")
        info.add_run(f"{self.ent_time.get()} час.\n")
        info.add_run(f"ул. Гоголя, 48, каб. {self.cabinet_number.get().strip() or '213'}")
        set_font(info, 14)

        doc.add_paragraph()

        # ===== Таблица =====
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False

        set_table_fixed_layout(table)
        set_table_width(table, 17.3)
        # № п/п делаем чуть шире, чтобы текст не переносился вертикально
        widths = [2, 7.10, 9.45]

        for i, width in enumerate(widths):
            set_cell_width(table.rows[0].cells[i], width)
        set_row_min_height(table.rows[0], 900)

        headers = ["№\nп/п", "ФИО", "Должность"]

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1

            run = p.add_run(header)
            run.bold = True
            run.font.name = "PT Astra Serif"
            run.font.size = Pt(14)

        def add_section_row(section_title):
            row = table.add_row()

            for col_idx, width in enumerate(widths):
                set_cell_width(row.cells[col_idx], width)

            merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            merged.text = section_title

            for paragraph in merged.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.line_spacing = 1

            set_cell_font(merged, 14)

            for paragraph in merged.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        def add_person_row(number, person_data):
            table_row = table.add_row()
            # Высота строки увеличена: после ФИО и должности добавляется
            # пустой абзац, чтобы между участниками был видимый интервал.
            set_row_min_height(table_row, 940)
            row = table_row.cells

            for col_idx, width in enumerate(widths):
                set_cell_width(row[col_idx], width)
                row[col_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            # № п/п
            row[0].text = f"{number}."
            center_cell_text(row[0])

            # ФИО: фамилия на первой строке, имя и отчество на второй,
            # затем пустой абзац для расстояния до следующей строки таблицы.
            row[1].text = ""
            surname, name_patronymic = format_fio_for_table(person_data["fio"])

            p_fio_1 = row[1].paragraphs[0]
            p_fio_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_fio_1.paragraph_format.line_spacing = 1
            p_fio_1.paragraph_format.space_before = Pt(0)
            p_fio_1.paragraph_format.space_after = Pt(0)
            p_fio_1.add_run(surname)

            if name_patronymic:
                p_fio_2 = row[1].add_paragraph()
                p_fio_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fio_2.paragraph_format.line_spacing = 1
                p_fio_2.paragraph_format.space_before = Pt(0)
                p_fio_2.paragraph_format.space_after = Pt(0)
                p_fio_2.add_run(name_patronymic)

            p_fio_empty = row[1].add_paragraph()
            p_fio_empty.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_fio_empty.paragraph_format.line_spacing = 1
            p_fio_empty.paragraph_format.space_before = Pt(0)
            p_fio_empty.paragraph_format.space_after = Pt(0)

            # Должность: после текста также добавляем пустой абзац,
            # чтобы правая ячейка была такой же высоты визуально.
            row[2].text = ""
            p_position = row[2].paragraphs[0]
            p_position.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_position.paragraph_format.line_spacing = 1
            p_position.paragraph_format.space_before = Pt(0)
            p_position.paragraph_format.space_after = Pt(0)
            p_position.add_run(person_data["dolzhnost"])

            p_position_empty = row[2].add_paragraph()
            p_position_empty.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_position_empty.paragraph_format.line_spacing = 1
            p_position_empty.paragraph_format.space_before = Pt(0)
            p_position_empty.paragraph_format.space_after = Pt(0)

            for cell in row:
                set_cell_font(cell, 14)

                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1

        row_number = 1

        def add_people_sorted_without_position_headers(main_section_title, people):
            """
            Добавляет только основной раздел:
            - Приглашенные
            - Органы местного самоуправления

            Внутренние заголовки по должностям НЕ выводятся.
            При этом сортировка сохраняется:
            сначала по должности от высшей к низшей,
            затем по алфавиту внутри каждой должности.
            """
            nonlocal row_number

            if not people:
                return

            add_section_row(main_section_title)

            sorted_people = self._sort_people_by_official_rank(people)

            for person in sorted_people:
                add_person_row(row_number, person)
                row_number += 1

        # Сначала приглашенные.
        add_people_sorted_without_position_headers(
            "Приглашенные:",
            invited_participants
        )

        # Потом органы МСУ.
        add_people_sorted_without_position_headers(
            "Органы местного самоуправления:",
            local_participants
        )

        # ===== Протокол ведет =====
        row = table.add_row()
        set_row_min_height(row, 520)

        for col_idx, width in enumerate(widths):
            set_cell_width(row.cells[col_idx], width)

        merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
        merged.text = "Протокол ведет:"
        justify_cell_text(merged)
        set_cell_font(merged, 14)

        table_row = table.add_row()
        set_row_min_height(table_row, 940)
        row = table_row.cells

        for col_idx, width in enumerate(widths):
            set_cell_width(row[col_idx], width)
            row[col_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        row[0].text = f"{row_number}."
        center_cell_text(row[0])

        keeper_name = self.protocol_keeper.get().strip() or "Иванова Елена Николаевна"
        surname, name_patronymic = format_fio_for_table(keeper_name)

        row[1].text = ""

        p_keeper_1 = row[1].paragraphs[0]
        p_keeper_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_keeper_1.paragraph_format.line_spacing = 1
        p_keeper_1.add_run(surname)

        if name_patronymic:
            p_keeper_2 = row[1].add_paragraph()
            p_keeper_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_keeper_2.paragraph_format.line_spacing = 1
            p_keeper_2.paragraph_format.space_before = Pt(0)
            p_keeper_2.paragraph_format.space_after = Pt(0)
            p_keeper_2.add_run(name_patronymic)

        p_keeper_empty = row[1].add_paragraph()
        p_keeper_empty.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_keeper_empty.paragraph_format.line_spacing = 1
        p_keeper_empty.paragraph_format.space_before = Pt(0)
        p_keeper_empty.paragraph_format.space_after = Pt(0)

        row[2].text = ""
        p_keeper_position = row[2].paragraphs[0]
        p_keeper_position.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_keeper_position.paragraph_format.line_spacing = 1
        p_keeper_position.paragraph_format.space_before = Pt(0)
        p_keeper_position.paragraph_format.space_after = Pt(0)
        p_keeper_position.add_run("главный специалист отдела судебной работы правового комитета")

        p_keeper_position_empty = row[2].add_paragraph()
        p_keeper_position_empty.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_keeper_position_empty.paragraph_format.line_spacing = 1
        p_keeper_position_empty.paragraph_format.space_before = Pt(0)
        p_keeper_position_empty.paragraph_format.space_after = Pt(0)

        for cell in row:
            set_cell_font(cell, 14)

        saved_path = self.safe_save_docx(doc, out_path)

        if not saved_path:
            return

        messagebox.showinfo("Успех", f"✅ Документ создан:\n{saved_path}")

    def _task_date_to_calendar_date(self, date_text):
        """
        Переводит дату из ДД.ММ.ГГГГ в ГГГГ-ММ-ДД.
        Именно такой формат использует kalendar.py.
        """
        date_text = (date_text or "").strip()

        if not date_text:
            return ""

        try:
            return datetime.strptime(date_text, "%d.%m.%Y").strftime("%Y-%m-%d")
        except Exception:
            return ""

    def _ensure_calendar_meeting_columns(self):
        """
        Добавляет связь календарной записи с совещанием.
        Нужна, чтобы из календаря по кнопке "Подробнее" открыть исходное совещание.
        """
        try:
            with get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("PRAGMA table_info(calendar_tasks)")
                existing_columns = [row[1] for row in cursor.fetchall()]

                if "meeting_id" not in existing_columns:
                    cursor.execute("ALTER TABLE calendar_tasks ADD COLUMN meeting_id INTEGER")

                conn.commit()
        except Exception as e:
            print(f"Ошибка добавления meeting_id в calendar_tasks: {e}")

    def transfer_unchecked_tasks_to_calendar(self, meeting_id=None):
        """
        Переносит в календарь только поручения без галочки.
        В календаре описание теперь единое:
        "Поручение из совещания по вопросу ...".
        Каждая календарная запись получает meeting_id, чтобы по кнопке
        "Подробнее" можно было открыть исходное совещание.
        """
        added_count = 0

        self._ensure_calendar_meeting_columns()

        meeting_title = self.ent_title.get().strip()
        meeting_time = self.ent_time.get().strip()
        calendar_description = f"Поручение из совещания по вопросу {meeting_title}"

        # Чтобы при повторном сохранении одного и того же совещания
        # календарь не наполнялся дублями, удаляем старые календарные
        # поручения, созданные именно из этого совещания.
        if meeting_id:
            try:
                with get_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        "DELETE FROM calendar_tasks WHERE meeting_id = ? AND task_type = ?",
                        (meeting_id, "поручение")
                    )
                    conn.commit()
            except Exception as e:
                print(f"Ошибка очистки старых календарных поручений совещания: {e}")

        for row in self.tasks_rows:
            task_text = row["task"].get().strip()
            task_date = row["date"].get().strip()
            task_done = row["done"].get()

            if not task_text:
                continue

            if task_done:
                continue

            calendar_date = self._task_date_to_calendar_date(task_date)

            if not calendar_date:
                continue

            calendar_task = {
                "date":calendar_date,
                "user":self.current_user,
                "type":"поручение",
                "desc":calendar_description,
                "time":meeting_time,
                "court":"",
                "subject":meeting_title,
                "assignment":"",
                "meeting_id":meeting_id
            }

            task_id = calendar_add_task(calendar_task)

            if meeting_id and task_id:
                try:
                    with get_connection() as conn:
                        cursor = conn.cursor()
                        cursor.execute(
                            "UPDATE calendar_tasks SET meeting_id = ? WHERE id = ?",
                            (meeting_id, task_id)
                        )
                        conn.commit()
                except Exception as e:
                    print(f"Ошибка сохранения meeting_id для календарной записи: {e}")

            added_count += 1

        return added_count

    def save_meeting(self):
        if not all([self.ent_title.get().strip(), self.ent_date.get().strip(), self.ent_time.get().strip()]):
            messagebox.showwarning("Ввод", "Заполните 'По вопросу', Дату и Время!")
            return

        tasks_data = []

        for i, t in enumerate(self.tasks_rows, start=1):
            committees = self._get_task_committees_from_row(t)
            committee_text = " §§ ".join(committees)
            task_text = t["task"].get().strip() if t.get("task") else ""
            date_text = t["date"].get().strip() if t.get("date") else ""
            done_text = "☑" if t["done"].get() else "☐"

            if committee_text:
                tasks_data.append(
                    f"{i}. {committee_text} | {task_text} | {date_text} | {done_text}"
                )
            else:
                tasks_data.append(
                    f"{i}. {task_text} | {date_text} | {done_text}"
                )

        try:
            if self.protocol_report_enabled.get():
                report_text = self.protocol_report_text.get().strip()
                report_date = self.protocol_report_date.get().strip()
                tasks_data.append(
                    f"__REPORT__ | {report_text} | {report_date} | ☑"
                )
        except Exception:
            pass

        data = {
            'title':self.ent_title.get().strip(),
            'date_val':self.ent_date.get().strip(),
            'time_val':self.ent_time.get().strip(),
            'agenda':self.get_agenda_text_for_save(),
            'protocol':"\n".join(tasks_data),
                        'participants': "\n".join(self.temp_selected_participants),
            'invited_participants': "\n".join(self.temp_invited_participants),
            'attachment_path':self.current_file_path,
            'has_participants_list':1 if self.attachment_type.get() == "participants" else 0,
            'has_telegram':1 if self.attachment_type.get() == "telegram" else 0
        }

        try:
            if self.current_id:
                meetings_update(self.current_id, data)
                saved_id = self.current_id
            else:
                new_id = meetings_add(data)
                saved_id = new_id if new_id else self._get_last_saved_meeting_id()
                self.current_id = saved_id

            self._save_meeting_extra_fields(saved_id)

            added_to_calendar = self.transfer_unchecked_tasks_to_calendar(saved_id)

            if added_to_calendar:
                messagebox.showinfo(
                    "Успех",
                    f"Сохранено!\nВ календарь перенесено поручений: {added_to_calendar}"
                )
            else:
                messagebox.showinfo("Успех", "Сохранено!")

            self.load_meetings()
            self.clear_form()

        except Exception as e:
            messagebox.showerror("БД", f"Сохранение: {e}")

    def delete_meeting(self):
        if not self.current_id:
            messagebox.showinfo("Инфо", "Выберите запись.")
            return
        if messagebox.askyesno("Подтверждение", "Удалить запись?"):
            try:
                meetings_delete(self.current_id)
                messagebox.showinfo("Успех", "Удалено.")
                self.clear_form()
                self.load_meetings()
            except Exception as e:
                messagebox.showerror("Ошибка", f"Удаление: {e}")

    def clear_form(self):
        self.current_id = None
        self.ent_title.delete(0, "end")
        self.ent_date.delete(0, "end");
        self.ent_date.insert(0, datetime.now().strftime("%d.%m.%Y"))
        self.ent_time.delete(0, "end");
        self.ent_time.insert(0, datetime.now().strftime("%H:%M"))
        self._set_agenda_questions_from_text("")
        self.txt_protocol.delete("1.0", "end")
        self.attachment_type.set("none")
        self.telegram_number.set("№ 200/05/ИТФ___")
        self.transfer_fio.set("Иванова Елена Николаевна")
        self.transfer_phone.set("")
        self.transfer_email.set("fedorova-en@barnaul-adm.ru")
        self.telegram_sign_fio.set("О.А. Финк")
        self.protocol_number.set("200/05/ПРОТ-___")
        self.protocol_chair_fio.set("О.А. Финк")
        self.protocol_chair_position.set("заместитель главы администрации города, руководитель аппарата")
        self.cabinet_number.set("213")
        self.protocol_keeper.set("Иванова Елена Николаевна")
        self.agenda_sign_position.set("Председатель правового комитета")
        self.agenda_sign_fio.set("О.И. Насыров")
        self.telegram_number_frame.pack_forget()

        for btn in self.doc_type_buttons.values():
            btn.configure(style="Modern.TButton")
        self.temp_selected_participants = []
        self.temp_invited_participants = []
        self.current_file_path = ""
        self.attached_files = []
        self.attachment_var.set("Файл не прикреплен")
        for i in self.tree.selection():
            self.tree.selection_remove(i)
        # Очищаем поручения и оставляем одну пустую строку
        self.tasks_rows.clear()
        self.add_task_row()
        try:
            self.protocol_report_enabled.set(False)
            self.protocol_report_text.set(
                "О проделанной работе проинформировать правовой комитет администрации города Барнаула до "
            )
            self.protocol_report_date.set("")
            if hasattr(self, "update_protocol_report_state"):
                self.update_protocol_report_state()
        except Exception:
            pass
        self.update_visible_fields_by_doc_type()


def _ask_initial_meeting_doc_type(parent):
    """
    Мини-окно выбора типа создаваемого документа.

    Важно: оно создается ДО основного окна MeetingsWindow. Так не возникает
    ситуации, когда основное окно скрыто через withdraw(), а модальный дочерний
    диалог с grab_set() оказывается недоступным и вся программа выглядит зависшей.
    """
    result = {"value": None}

    dialog = tk.Toplevel(parent)
    dialog.title("Выбор документа")
    dialog.configure(bg="#f4f6f9")
    dialog.resizable(False, False)

    try:
        dialog.transient(parent)
        dialog.grab_set()
        apply_modern_style(dialog)
    except Exception:
        pass

    container = tk.Frame(dialog, bg="#f4f6f9")
    container.pack(fill="both", expand=True, padx=22, pady=18)

    tk.Label(
        container,
        text="Выберите тип создаваемого документа",
        font=("Segoe UI", 13, "bold"),
        bg="#f4f6f9",
        fg="#111111"
    ).pack(anchor="center", pady=(0, 16))

    buttons_frame = tk.Frame(container, bg="#f4f6f9")
    buttons_frame.pack(fill="both", expand=True)

    def choose(doc_type):
        result["value"] = doc_type
        try:
            dialog.grab_release()
        except Exception:
            pass
        dialog.destroy()

    for title, doc_type in (
        ("Список участников", "participants"),
        ("Повестка", "agenda"),
        ("Телефонограмма", "telegram"),
        ("Протокол", "protocol"),
    ):
        ttk.Button(
            buttons_frame,
            text=title,
            style="Modern.TButton",
            command=lambda value=doc_type: choose(value)
        ).pack(fill="x", pady=5, ipady=4)

    def close_without_choice():
        result["value"] = None
        try:
            dialog.grab_release()
        except Exception:
            pass
        dialog.destroy()

    dialog.protocol("WM_DELETE_WINDOW", close_without_choice)
    dialog.bind("<Escape>", lambda event: close_without_choice())

    dialog.update_idletasks()
    width = max(dialog.winfo_width(), 430)
    height = max(dialog.winfo_height(), 260)
    screen_w = dialog.winfo_screenwidth()
    screen_h = dialog.winfo_screenheight()
    x = (screen_w - width) // 2
    y = (screen_h - height) // 2
    dialog.geometry(f"{width}x{height}+{x}+{y}")
    dialog.lift()
    dialog.focus_force()

    parent.wait_window(dialog)
    return result["value"]


def show_meetings(parent, current_user=None):
    initial_doc_type = _ask_initial_meeting_doc_type(parent)

    if not initial_doc_type:
        return

    win = MeetingsWindow(parent, current_user, initial_doc_type=initial_doc_type)
    win.transient(parent)
    parent.wait_window(win)